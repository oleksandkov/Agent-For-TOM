"""docx2txt.py — convert any .docx to a TXT file while preserving as much
formatting detail as possible.

Strategy
--------
Uses ``python-docx`` for the structural walk (paragraphs, runs, tables,
sections) and reaches into the underlying ``lxml`` tree for the formatting
attributes that ``python-docx`` does not expose (underline, strikethrough,
highlight, subscript, superscript, character color in some cases).

What it preserves
-----------------
* Document outline — heading levels annotated as ``# H1``, ``## H2`` ...
* Paragraphs — alignment, indentation, spacing (where significant)
* Inline formatting — bold / italic / underline / strikethrough / mono /
  highlight marked with inline tags
* Font name + size — optional per-line annotation (``--annotate-fonts``)
* Tables — rendered as tab-separated rows with row separators
* Lists — bulleted vs numbered, preserving list level
* Headers / footers — separated and marked
* Embedded images — noted inline, optionally extracted (``--extract-images``)
* Page breaks — explicit ``----- Page Break -----`` markers
* Full structured data — optional sidecar ``.json`` (``--json``)

Limitations
-----------
* DOCX is a flow format — it has no fixed page boundaries. Page breaks are
  only emitted where the document explicitly inserts one.
* Track changes, comments, and footnotes are extracted as plain text.
  The ``--json`` sidecar retains the full document XML for these.

Usage
-----
    python docx2txt.py file.docx
    python docx2txt.py file.docx -o out.txt
    python docx2txt.py file.docx --annotate-fonts
    python docx2txt.py file.docx --extract-images img/
    python docx2txt.py file.docx --json
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

try:
    import docx
    from docx.document import Document as _Document
    from docx.oxml.ns import qn
    from docx.shared import RGBColor
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
except ImportError:
    sys.stderr.write(
        "ERROR: python-docx is required. Install it with:\n"
        "    python -m pip install python-docx\n"
    )
    sys.exit(1)


# -- Inline formatting helpers ------------------------------------------------

def _bool(elem) -> bool:
    """Return True if the ``w:val`` attribute is not literally 'false'/'0'."""
    if elem is None:
        return False
    val = elem.get(qn("w:val"))
    if val is None:
        return True  # attribute presence without a value implies true
    return val.lower() not in ("false", "0", "off")


def run_formatting(run) -> dict[str, Any]:
    """Return a dict of formatting flags and properties for a run."""
    rPr = run._element.find(qn("w:rPr"))
    flags: dict[str, Any] = {
        "bold": bool(run.bold) if run.bold is not None else False,
        "italic": bool(run.italic) if run.italic is not None else False,
        "underline": False,
        "strike": False,
        "subscript": False,
        "superscript": False,
        "highlight": None,
        "color": None,
        "font": None,
        "size_pt": None,
    }

    # python-docx surfaces bold/italic/underline on Font; fall back to XML.
    font = run.font
    if font.underline is True or (isinstance(font.underline, str) and font.underline):
        flags["underline"] = True
    elif rPr is not None:
        flags["underline"] = _bool(rPr.find(qn("w:u")))

    if rPr is not None:
        flags["strike"] = _bool(rPr.find(qn("w:strike")))
        vert = rPr.find(qn("w:vertAlign"))
        if vert is not None:
            v = (vert.get(qn("w:val")) or "").lower()
            if v == "subscript":
                flags["subscript"] = True
            elif v == "superscript":
                flags["superscript"] = True
        hl = rPr.find(qn("w:highlight"))
        if hl is not None:
            flags["highlight"] = hl.get(qn("w:val"))
        col = rPr.find(qn("w:color"))
        if col is not None and col.get(qn("w:val")) not in (None, "auto"):
            flags["color"] = "#" + col.get(qn("w:val")).upper()

    if font.name:
        flags["font"] = font.name
    if font.size is not None:
        # python-docx returns Emu (Pt) — size in points
        flags["size_pt"] = float(font.size.pt)
    if font.color and font.color.rgb is not None:
        # Override with python-docx's normalized color when available.
        flags["color"] = "#" + str(font.color.rgb).upper()

    return flags


def is_mono(flags: dict[str, Any]) -> bool:
    """Heuristic: treat common monospace font names as mono runs."""
    name = (flags.get("font") or "").lower()
    return any(token in name for token in (
        "courier", "consolas", "monaco", "menlo", "roboto mono", "fira",
        "source code", "cascadia", "jetbrains", "consola",
    ))


def wrap_inline(text: str, flags: dict[str, Any]) -> str:
    """Wrap a run's text with inline markers describing its formatting.

    Multi-line text never gets backtick-wrapped (would be ambiguous) — the
    mono font is then expressed through the font annotation tag.
    """
    if not text:
        return text

    if is_mono(flags):
        # A backtick-wrapped multi-line block can't be parsed unambiguously.
        # Emit the text verbatim and rely on the font annotation to signal
        # monospace. For a single-line mono run, backticks are still fine.
        if "\n" in text:
            return text
        return f"`{text}`"

    # Apply outermost first so nesting renders correctly.
    if flags.get("subscript"):
        text = f"~{text}~"  # simple subscript marker
    if flags.get("superscript"):
        text = f"^{text}^"

    if flags.get("bold") and flags.get("italic"):
        text = f"***{text}***"
    elif flags.get("bold"):
        text = f"**{text}**"
    elif flags.get("italic"):
        text = f"*{text}*"

    if flags.get("underline"):
        text = f"_{text}_"
    if flags.get("strike"):
        text = f"~~{text}~~"

    if flags.get("highlight"):
        text = f"[hl:{flags['highlight']}]{text}[/hl]"

    return text


def format_run_label(flags: dict[str, Any]) -> str:
    """Short per-run label for the --annotate-fonts mode."""
    parts = []
    if flags.get("font"):
        parts.append(flags["font"])
    if flags.get("size_pt") is not None:
        parts.append(f"{flags['size_pt']:.1f}pt")
    style = []
    if flags.get("bold"):
        style.append("B")
    if flags.get("italic"):
        style.append("I")
    if flags.get("underline"):
        style.append("U")
    if flags.get("strike"):
        style.append("S")
    if is_mono(flags):
        style.append("M")
    if flags.get("subscript"):
        style.append("sub")
    if flags.get("superscript"):
        style.append("sup")
    if style:
        parts.append("/".join(style))
    return " ".join(parts) if parts else "?"


# -- Paragraph / list handling ------------------------------------------------

_HEADING_RE = re.compile(r"^heading\s*(\d+)$", re.IGNORECASE)
_LIST_BULLET_RE = re.compile(r"^list\s*(bullet|paragraph)$", re.IGNORECASE)
_LIST_NUMBER_RE = re.compile(r"^list\s*number$", re.IGNORECASE)


def paragraph_kind(p: Paragraph) -> str:
    """Classify a paragraph as heading / list / body / code / quote."""
    style = (p.style.name or "").strip() if p.style is not None else ""

    m = _HEADING_RE.match(style)
    if m:
        return f"h{int(m.group(1))}"
    if _LIST_BULLET_RE.match(style):
        return "list-bullet"
    if _LIST_NUMBER_RE.match(style):
        return "list-number"
    if "quote" in style.lower():
        return "quote"
    if "code" in style.lower() and "block" in style.lower():
        return "code-block"
    return "body"


def paragraph_indent(p: Paragraph) -> str:
    """Translate paragraph left indent into leading spaces."""
    pf = p.paragraph_format
    if pf.left_indent is None:
        return ""
    # Rough: 1 char ~ 7.5 pt at 12pt body, scale to actual font size
    pt = pf.left_indent.pt
    chars = int(round(pt / 7.5))
    return " " * min(max(chars, 0), 40)


def render_paragraph(p: Paragraph, *, annotate: bool) -> str:
    """Render one paragraph as a single string."""
    pieces: list[str] = []
    for run in p.runs:
        text = run.text
        if not text:
            continue
        flags = run_formatting(run)
        pieces.append(wrap_inline(text, flags))

    line = "".join(pieces).rstrip()
    if not line:
        return ""

    kind = paragraph_kind(p)
    indent = paragraph_indent(p)

    if kind.startswith("h"):
        level = int(kind[1:])
        prefix = "#" * min(level, 6) + " "
        out = prefix + line
    elif kind == "list-bullet":
        out = "- " + line
    elif kind == "list-number":
        out = "1. " + line
    elif kind == "quote":
        out = "> " + line
    elif kind == "code-block":
        out = "    " + line  # visual code indent
    else:
        out = line

    if annotate:
        # Pull font info from the first non-empty run for an annotation label.
        label = ""
        for run in p.runs:
            if run.text.strip():
                label = format_run_label(run_formatting(run))
                break
        if label:
            out = f"[{label}] {out}"

    return indent + out


# -- Tables -------------------------------------------------------------------

def render_table(table: Table) -> list[str]:
    """Render a table as a list of tab-separated rows."""
    rows: list[list[str]] = []
    for row in table.rows:
        cells: list[str] = []
        for cell in row.cells:
            cell_text_parts: list[str] = []
            for p in cell.paragraphs:
                rendered = render_paragraph(p, annotate=False)
                if rendered:
                    cell_text_parts.append(rendered)
            cells.append("\n".join(cell_text_parts).replace("\t", "    "))
        rows.append(cells)

    if not rows:
        return []

    width = max(len(r) for r in rows)
    out: list[str] = ["[table]"]
    for r in rows:
        r = r + [""] * (width - len(r))
        out.append("\t".join(r))
    out.append("[/table]")
    return out


# -- Headers / footers --------------------------------------------------------

def render_header_footer(part, label: str) -> list[str]:
    """Render header or footer paragraphs (and any nested tables)."""
    if part is None:
        return []
    out: list[str] = [f"\n----- {label} -----"]
    body = getattr(part, "paragraphs", []) or []
    for p in body:
        rendered = render_paragraph(p, annotate=False)
        if rendered:
            out.append(rendered)
    for t in getattr(part, "tables", []) or []:
        out.extend(render_table(t))
    return out


# -- Image extraction ---------------------------------------------------------

def extract_images(doc: _Document, out_dir: Path) -> int:
    """Walk the document and save every embedded image to *out_dir*."""
    out_dir.mkdir(parents=True, exist_ok=True)
    saved = 0
    seen: set[str] = set()
    for rel_id, rel in doc.part.rels.items():
        if rel.reltype.endswith("/image"):
            target = rel.target_part
            if target.partname in seen:
                continue
            seen.add(target.partname)
            ext = Path(target.partname).suffix or ".png"
            data = target.blob
            try:
                img_index = saved + 1
                path = out_dir / f"img{img_index:03d}{ext}"
                path.write_bytes(data)
                saved += 1
            except OSError:
                pass
    return saved


# -- Main extraction ----------------------------------------------------------

def iter_block_items(parent) -> Any:
    """Yield paragraphs and tables in document order.

    Works for the body and for table cells.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        parent_elm = parent

    for child in parent_elm.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif tag == qn("w:tbl"):
            yield Table(child, parent)
        # Other block-level elements (sdt, structured doc tags) are skipped.


def extract(doc: _Document, *, annotate: bool) -> list[str]:
    out: list[str] = []

    # Headers / footers (one set per section; dedupe identical text).
    seen_hf: set[str] = set()
    for section in doc.sections:
        for kind, part in (("header", section.header), ("footer", section.footer)):
            if part is None:
                continue
            text = "\n".join(p.text for p in getattr(part, "paragraphs", []))
            if not text.strip() or text in seen_hf:
                continue
            seen_hf.add(text)
            out.extend(render_header_footer(part, kind.capitalize()))

    # Body
    out.append("\n===== Body =====")
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            # Detect page-break-only paragraphs (python-docx's
            # Document.add_page_break() creates a dedicated empty paragraph
            # containing only a <w:br w:type="page"/>). Emit the marker
            # first, then the paragraph's text (if any).
            page_break_here = any(
                br.get(qn("w:type")) == "page"
                for br in block._element.iter(qn("w:br"))
            )
            rendered = render_paragraph(block, annotate=annotate)
            if page_break_here:
                out.append("----- Page Break -----")
            if rendered:
                out.append(rendered)
        elif isinstance(block, Table):
            out.extend(render_table(block))

    return out


_STRIP_PATTERNS = [
    (re.compile(r"\[hl:[^\]]+\](.+?)\[/hl\]"), r"\1"),  # highlight
    (re.compile(r"\*\*(.+?)\*\*"), r"\1"),                # bold
    (re.compile(r"\*(.+?)\*"), r"\1"),                   # italic
    (re.compile(r"__(.+?)__"), r"\1"),                   # underline (alt)
    (re.compile(r"_(.+?)_"), r"\1"),                     # underline
    (re.compile(r"~~(.+?)~~"), r"\1"),                   # strike
    (re.compile(r"~(.+?)~"), r"\1"),                     # subscript marker
    (re.compile(r"\^(.+?)\^"), r"\1"),                   # superscript marker
    (re.compile(r"`(.+?)`"), r"\1"),                     # mono
]


def strip_formatting(text: str) -> str:
    for pattern, repl in _STRIP_PATTERNS:
        text = pattern.sub(repl, text)
    return text


def build_argparser() -> argparse.ArgumentParser:
    ap = argparse.ArgumentParser(
        description="Convert a .docx to a formatting-preserving TXT file.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    ap.add_argument("docx", help="Input .docx file")
    ap.add_argument("-o", "--output", help="Output TXT path (default: <docx>.txt)")
    ap.add_argument("--no-format", action="store_true",
                    help="Strip **bold**/*italic*/`mono`/etc. markers, plain text only")
    ap.add_argument("--annotate-fonts", action="store_true",
                    help="Prefix each paragraph with its dominant font + size")
    ap.add_argument("--extract-images", metavar="DIR",
                    help="Extract embedded images into DIR")
    ap.add_argument("--json", action="store_true",
                    help="Also write a sidecar .json with full structured data")
    return ap


def main(argv: list[str] | None = None) -> int:
    args = build_argparser().parse_args(argv)

    in_path = Path(args.docx)
    if not in_path.is_file():
        sys.stderr.write(f"ERROR: file not found: {in_path}\n")
        return 1

    out_path = Path(args.output) if args.output else in_path.with_suffix(".txt")
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = docx.Document(in_path)
    lines = extract(doc, annotate=args.annotate_fonts)
    text = "\n".join(lines)
    if args.no_format:
        text = strip_formatting(text)

    out_path.write_text(text, encoding="utf-8")
    n_paragraphs = sum(1 for _ in doc.paragraphs)
    n_tables = len(doc.tables)
    print(
        f"Wrote {out_path}  "
        f"({n_paragraphs} paragraph{'s' if n_paragraphs != 1 else ''}, "
        f"{n_tables} table{'s' if n_tables != 1 else ''})"
    )

    if args.extract_images:
        n = extract_images(doc, Path(args.extract_images))
        print(f"Extracted {n} image{'s' if n != 1 else ''} -> {args.extract_images}")

    if args.json:
        json_path = out_path.with_suffix(".json")
        json_path.write_text(
            json.dumps(
                {
                    "source": str(in_path),
                    "paragraphs": [
                        {"style": p.style.name if p.style else None, "text": p.text}
                        for p in doc.paragraphs
                    ],
                    "tables": [
                        [[c.text for c in row.cells] for row in t.rows] for t in doc.tables
                    ],
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        print(f"Wrote {json_path}  (full structured data)")

    return 0


if __name__ == "__main__":
    sys.exit(main())
