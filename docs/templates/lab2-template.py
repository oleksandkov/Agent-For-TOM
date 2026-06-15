import os
import docx
from docx.shared import Pt
from docx.shared import Cm as DocxCm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.units import cm as PdfCm

def create_docx(filename):
    """Генерує DOCX файл зі структурою Лабораторної роботи №2 (з Варіантами)"""
    doc = docx.Document()
    
    # Налаштування стилю "Normal" для всього документа
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    # Налаштування інтервалу (1.5) та вирівнювання за замовчуванням
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.first_line_indent = DocxCm(1.25) # Червоний рядок
    paragraph_format.space_after = Pt(0)

    # 1. Заголовок "ЛАБОРАТОРНА РОБОТА"
    p_header = doc.add_paragraph()
    p_header.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header.paragraph_format.first_line_indent = DocxCm(0)
    p_header.add_run("ЛАБОРАТОРНА РОБОТА № [ВСТАВТЕ НОМЕР]")

    # 2. Назва роботи
    p_title = doc.add_paragraph()
    p_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.first_line_indent = DocxCm(0)
    run_title = p_title.add_run("[Вставте назву лабораторної роботи]")
    run_title.bold = True
    
    doc.add_paragraph()

    # 3. Мета роботи
    p_meta = doc.add_paragraph()
    run_meta = p_meta.add_run("Мета роботи")
    run_meta.bold = True
    p_meta.add_run(" - [вставте текст мети роботи]")

    doc.add_paragraph()

    # 4. Загальні відомості
    p_zagalni = doc.add_paragraph()
    p_zagalni.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_zagalni.paragraph_format.first_line_indent = DocxCm(0)
    p_zagalni.add_run("Загальні відомості").bold = True

    doc.add_paragraph("[Вставте теоретичний матеріал тут. Форматування налаштовано автоматично.]")
    
    doc.add_paragraph()

    # 5. Контрольні запитання (У Лаб №2 вони йдуть ПЕРЕД завданням)
    p_kontr = doc.add_paragraph()
    p_kontr.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_kontr.paragraph_format.first_line_indent = DocxCm(0)
    p_kontr.add_run("Контрольні запитання").bold = True

    doc.add_paragraph("1. [Вставте перше запитання]")
    doc.add_paragraph("2. [Вставте друге запитання]")

    doc.add_paragraph()

    # 6. Завдання.
    p_zavd = doc.add_paragraph()
    p_zavd.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_zavd.paragraph_format.first_line_indent = DocxCm(0)
    p_zavd.add_run("Завдання.").bold = True

    doc.add_paragraph("1. [Вставте перше завдання]")
    doc.add_paragraph("2. [Вставте друге завдання]")

    doc.add_paragraph()

    # 7. Варіанти (З абзацного відступу, жирним)
    p_var = doc.add_paragraph()
    p_var.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_var = p_var.add_run("Варіанти")
    run_var.bold = True

    # Список варіантів
    doc.add_paragraph("1. [Значення варіанту 1]")
    doc.add_paragraph("2. [Значення варіанту 2]")
    doc.add_paragraph("3. [Значення варіанту 3]")

    doc.add_paragraph()

    # 8. Література
    p_lit = doc.add_paragraph()
    p_lit.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_lit.paragraph_format.first_line_indent = DocxCm(0)
    p_lit.add_run("Література").bold = True

    doc.add_paragraph("1. [Вставте перше джерело]")
    doc.add_paragraph("2. [Вставте друге джерело]")

    doc.save(filename)
    print(f"Успішно створено файл: {filename}")


def create_pdf(filename):
    """Генерує PDF зі структурою Лабораторної роботи №2 (з Варіантами)"""
    
    try:
        pdfmetrics.registerFont(TTFont('TimesNewRoman', 'times.ttf'))
        pdfmetrics.registerFont(TTFont('TimesNewRoman-Bold', 'timesbd.ttf'))
        font_regular = 'TimesNewRoman'
        font_bold = 'TimesNewRoman-Bold'
    except Exception as e:
        print("Шрифти Times New Roman не знайдено, використовується стандартний.")
        font_regular = 'Helvetica'
        font_bold = 'Helvetica-Bold'

    doc = SimpleDocTemplate(filename, pagesize=A4,
                            rightMargin=1.5*PdfCm, leftMargin=3.0*PdfCm,
                            topMargin=2.0*PdfCm, bottomMargin=2.0*PdfCm)
    
    # Стилі
    style_center = ParagraphStyle(
        name='CenterNormal', fontName=font_regular, fontSize=14, leading=21, 
        alignment=TA_CENTER, spaceAfter=14)
    
    style_center_bold = ParagraphStyle(
        name='CenterBold', fontName=font_bold, fontSize=14, leading=21, 
        alignment=TA_CENTER, spaceAfter=14, spaceBefore=14)
    
    style_body = ParagraphStyle(
        name='BodyJustify', fontName=font_regular, fontSize=14, leading=21, 
        alignment=TA_JUSTIFY, firstLineIndent=35, spaceAfter=0) # 35 point ~ 1.25 cm

    style_body_bold = ParagraphStyle(
        name='BodyBold', fontName=font_bold, fontSize=14, leading=21, 
        alignment=TA_LEFT, firstLineIndent=35, spaceAfter=10, spaceBefore=10)

    story = []

    # 1. Заголовок
    story.append(Paragraph("ЛАБОРАТОРНА РОБОТА № [ВСТАВТЕ НОМЕР]", style_center))
    # 2. Назва
    story.append(Paragraph("[Вставте назву лабораторної роботи]", style_center_bold))
    
    # 3. Мета
    story.append(Paragraph("<b>Мета роботи</b> - [вставте текст мети роботи]", style_body))
    story.append(Spacer(1, 14))
    
    # 4. Загальні відомості
    story.append(Paragraph("Загальні відомості", style_center_bold))
    story.append(Paragraph("[Вставте теоретичний матеріал тут. Форматування налаштовано автоматично.]", style_body))
    
    # 5. Контрольні запитання
    story.append(Paragraph("Контрольні запитання", style_center_bold))
    story.append(Paragraph("1. [Вставте перше запитання]", style_body))
    story.append(Paragraph("2. [Вставте друге запитання]", style_body))

    # 6. Завдання
    story.append(Paragraph("Завдання.", style_center_bold))
    story.append(Paragraph("1. [Вставте перше завдання]", style_body))
    story.append(Paragraph("2. [Вставте друге завдання]", style_body))
    
    # 7. Варіанти
    story.append(Paragraph("Варіанти", style_body_bold))
    story.append(Paragraph("1. [Значення варіанту 1]", style_body))
    story.append(Paragraph("2. [Значення варіанту 2]", style_body))
    story.append(Paragraph("3. [Значення варіанту 3]", style_body))

    # 8. Література
    story.append(Paragraph("Література", style_center_bold))
    story.append(Paragraph("1. [Вставте перше джерело]", style_body))
    story.append(Paragraph("2. [Вставте друге джерело]", style_body))

    doc.build(story)
    print(f"Успішно створено файл: {filename}")

if __name__ == "__main__":
    # Отримуємо точний абсолютний шлях до папки, де лежить цей скрипт
    current_dir = os.path.dirname(os.path.abspath(__file__))
    
    # Створюємо повні шляхи для збереження нових файлів
    docx_path = os.path.join(current_dir, "Lab_Template_Lab2_Style.docx")
    pdf_path = os.path.join(current_dir, "Lab_Template_Lab2_Style.pdf")

    # Запускаємо генерацію
    create_docx(docx_path)
    create_pdf(pdf_path)