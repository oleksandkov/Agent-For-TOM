import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH as A
from docx.shared import Pt, Cm

OUT = r'labwork_sessions\nemotron-3.5-lightning-free-session\AI_Integration_LabWork.docx'
FONT = 'Times New Roman'
SIZE_TITLE = Pt(20)   # Title is 20pt in original
SIZE_BODY = Pt(14)    # Body text is 14pt

d = docx.Document()
d.styles['Normal'].font.name = FONT
d.styles['Normal'].font.size = SIZE_BODY

def para(text='', *, align=A.JUSTIFY, bold=False, italic=False,
         indent=None, style=None):
    p = d.add_paragraph(style=style)
    p.alignment = align
    if indent is not None:
        p.paragraph_format.first_line_indent = indent
    if text:
        r = p.add_run(text)
        r.bold, r.italic = bold, italic
        r.font.name, r.font.size = FONT, SIZE_BODY
    return p

def centre(t):
    p = para(t, align=A.CENTER, bold=True)
    if p.runs:
        p.runs[0].font.size = SIZE_TITLE
    return p

def body(t):
    p = para(t, align=A.JUSTIFY, indent=Cm(1.25))
    if p.runs:
        p.runs[0].font.size = SIZE_BODY
    return p

def item(t):
    p = para(t, align=A.JUSTIFY, style='List Number')
    if p.runs:
        p.runs[0].font.size = SIZE_BODY
    return p

def lead(label, rest):
    p = para(align=A.JUSTIFY, indent=Cm(1.25))
    for txt, b in ((label, True), (rest, False)):
        if p.runs:
            p.runs[-1].bold = b
        else:
            r = p.add_run(txt); r.bold = b
            r.font.name, r.font.size = FONT, SIZE_BODY
    return p

# ---- Title page ----
for _ in range(8):
    para()

centre('МЕТОДИЧНІ ВКАЗІВКИ ДО ВИКОНАННЯ')
centre('ЛАБОРАТОРНИХ РОБОТ З ДИСЦИПЛІНИ')
centre('«ІНТЕГРАЦІЯ ШТУЦЬКОВОГО ИНТЕЛЛЕКТУ В ЗАстосунки»')
d.add_page_break()

# ---- Lab 1 header ----
centre('ЛАБОРАТОРНА РОБОТА №1')
centre('Інтеграція штучного інтелекту в застосунки')
centre('')
lead('Мета роботи', ' - вивчити особливості інтеграції штучного інтелекту в програмні системи.')
centre('Загальні відомості')
body('Штучний інтелект (AI) increasingly integrates into various applications, providing intelligent features such as natural language processing, computer vision, predictive analytics, and decision support systems. This laboratory work explores the integration of AI capabilities into software applications, covering the fundamentals of AI integration, popular frameworks, and practical implementation approaches.')

# VNTU details block
p = para(align=A.JUSTIFY, indent=Cm(1.25))
r = p.add_run('Факультет: '); r.bold = True; r.font.name, r.font.size = FONT, SIZE_BODY
r = p.add_run('Факультет інформаційних технологій та комп\'ютерної інженерії'); r.font.name, r.font.size = FONT, SIZE_BODY

r = p.add_run('; Кафедра: '); r.bold = True; r.font.name, r.font.size = FONT, SIZE_BODY
r = p.add_run('Кафедра комп\'ютерних наук'); r.font.name, r.font.size = FONT, SIZE_BODY

r = p.add_run('; Спеціальність: '); r.bold = True; r.font.name, r.font.size = FONT, SIZE_BODY
r = p.add_run('122 Комп\'ютерні науки'); r.font.name, r.font.size = FONT, SIZE_BODY

r = p.add_run('; Місто: '); r.bold = True; r.font.name, r.font.size = FONT, SIZE_BODY
r = p.add_run('Вінниця'); r.font.name, r.font.size = FONT, SIZE_BODY

centre('Контрольні запитання')
item('1. Какі основні способи інтеграції штучного інтелекту в існуючі застосунки ви знаєте?')
item('2. Які переваги надає інтеграція AI в програмні системи?')
item('3. Які популярні фреймворки та бібліотеки використовуються для інтеграції AI?')
item('4. Які виникають при інтеграції AI виклики та як їх вирішувати?')
item('5. Як оцінювати ефективність інтегрованих AI-розв\'язків?')

centre('Завдання')
item('1. Розробити прототип застосунку з базовою інтеграцією штучного інтелекту (наприклад, чат-бот або система рекомендацій).')
item('2. Інтегрувати бібліотеку обробки мови природного (NLP) у консольний застосунок для аналізу тексту.')
item('3. Створити просту модель машинного навчання та інтегрувати її у веб-застосунок для передбачення цін.')
item('4. Реалізувати механізм випереджального введення відповідей для пришвидшення взаємодії з AI-системою.')
item('5. Оцінити продуктивність та ефективність інтегрованого AI-розв\'язку на тестовому наборі даних.')

centre('Зауваження. Допускається розробка та реалізація самостійного завдання студентом по узгодженню з викладачем.')
centre('Література')
item('1. Человеко-машинне взаимодействие: Учебне пособие / Новосибирский государственный технический университет. – Новосибирск, 2006. – 97 с.')
item('2. Інтерфейс "Користувач - комп\'ютер": Навчальний посібник / В.П. Майданюк, А.М. Петюх. – Вінниця: ВДТУ, 1999. – 66 с.')
item('3. Мандел Т. Разработка пользовательского интерфейса / Т. Мандел. – Пер. с англ. – М.: ДМК Пресс, 2001. – 416 с.')
item('4. Коутс Р. Інтерфейс "Человек - компьютер" / Р. Коутс, И. Влейминк. – Пер. с англ. – М.: Мир, 1990. – 501 с.')
item('5. Человеко-машинне взаимодействие: Учебне пособие / Новосибирський державний технічний університет. – Новосибирськ, 2006. – 97 с.')

d.save(OUT)
print('saved', OUT)