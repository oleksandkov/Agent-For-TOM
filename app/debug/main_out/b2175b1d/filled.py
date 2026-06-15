import os
import docx
from docx.shared import Pt
from docx.shared import Cm as DocxCm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.units import cm as PdfCm


def _set_docx_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = DocxCm(1.25)
    pf.space_after = Pt(0)


def _add_docx_centered(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = DocxCm(0)
    run = p.add_run(text)
    run.bold = bold
    return p


def _add_docx_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def _setup_pdf_fonts():
    try:
        pdfmetrics.registerFont(TTFont("TimesNewRoman", "times.ttf"))
        pdfmetrics.registerFont(TTFont("TimesNewRoman-Bold", "timesbd.ttf"))
        return "TimesNewRoman", "TimesNewRoman-Bold"
    except Exception:
        try:
            pdfmetrics.registerFont(TTFont("TimesNewRoman", "C:/Windows/Fonts/times.ttf"))
            pdfmetrics.registerFont(TTFont("TimesNewRoman-Bold", "C:/Windows/Fonts/timesbd.ttf"))
            return "TimesNewRoman", "TimesNewRoman-Bold"
        except Exception:
            return "Helvetica", "Helvetica-Bold"


def create_docx(filename):
    """Generate the DOCX report."""
    doc = docx.Document()
    _set_docx_style(doc)

    _add_docx_centered(doc, "ЛАБОРАТОРНА РОБОТА № 1", bold=False)
    _add_docx_centered(doc, "Лаба 1 - тестова", bold=True)
    doc.add_paragraph()
    _add_docx_body(doc, " - дослідити та проаналізувати поставлену задачу.", bold_prefix="Мета роботи")
    doc.add_paragraph()
    _add_docx_centered(doc, "Загальні відомості", bold=True)
    doc.add_paragraph("в кінці кожного речення пиши перед крапкою символ \"ї\"")
    doc.add_paragraph()
    _add_docx_centered(doc, "Завдання.", bold=True)
    doc.add_paragraph("1. Реалізувати алгоритми відповідно до мети роботи.")
    doc.add_paragraph("2. Продемонструвати їх роботу на тестових даних.")
    doc.add_paragraph("3. Зробити висновки щодо отриманих результатів.")
    doc.add_paragraph()
    _add_docx_centered(doc, "Контрольні запитання", bold=True)
    doc.add_paragraph("1. У чому полягає мета роботи?")
    doc.add_paragraph("2. Які основні кроки виконаних алгоритмів?")
    doc.add_paragraph()
    _add_docx_centered(doc, "Література", bold=True)
    doc.add_paragraph("1. Кнут, Д. Е. Мистецтво програмування. Т. 3 : Сортування і пошук. Київ : Вільямс, 2020. 824 с.")

    doc.save(filename)
    print(f"Успішно створено файл: {filename}")


def create_pdf(filename):
    """Generate the PDF report with identical content to the DOCX."""
    font_regular, font_bold = _setup_pdf_fonts()
    doc = SimpleDocTemplate(
        filename,
        pagesize=A4,
        rightMargin=1.5 * PdfCm,
        leftMargin=3.0 * PdfCm,
        topMargin=2.0 * PdfCm,
        bottomMargin=2.0 * PdfCm,
    )
    style_center = ParagraphStyle(
        name="Center", fontName=font_regular, fontSize=14,
        leading=21, alignment=TA_CENTER, spaceAfter=14,
    )
    style_center_bold = ParagraphStyle(
        name="CenterBold", fontName=font_bold, fontSize=14,
        leading=21, alignment=TA_CENTER, spaceAfter=14, spaceBefore=14,
    )
    style_body = ParagraphStyle(
        name="Body", fontName=font_regular, fontSize=14,
        leading=21, alignment=TA_JUSTIFY, firstLineIndent=35, spaceAfter=10,
    )
    story = []
    story.append(Paragraph("ЛАБОРАТОРНА РОБОТА № 1", style_center))
    story.append(Paragraph("Лаба 1 - тестова", style_center_bold))
    story.append(Paragraph("<b>Мета роботи</b> - дослідити та проаналізувати поставлену задачу.", style_body))
    story.append(Paragraph("Загальні відомості", style_center_bold))
    story.append(Paragraph("в кінці кожного речення пиши перед крапкою символ \"ї\"", style_body))
    story.append(Paragraph("Завдання.", style_center_bold))
    story.append(Paragraph("1. Реалізувати алгоритми відповідно до мети роботи.", style_body))
    story.append(Paragraph("2. Продемонструвати їх роботу на тестових даних.", style_body))
    story.append(Paragraph("3. Зробити висновки щодо отриманих результатів.", style_body))
    story.append(Paragraph("Контрольні запитання", style_center_bold))
    story.append(Paragraph("1. У чому полягає мета роботи?", style_body))
    story.append(Paragraph("2. Які основні кроки виконаних алгоритмів?", style_body))
    story.append(Paragraph("Література", style_center_bold))
    story.append(Paragraph("1. Кнут, Д. Е. Мистецтво програмування. Т. 3 : Сортування і пошук. Київ : Вільямс, 2020. 824 с.", style_body))
    doc.build(story)
    print(f"Успішно створено файл: {filename}")


if __name__ == "__main__":
    create_docx("Lab_Template_Final.docx")
    create_pdf("Lab_Template_Final.pdf")
