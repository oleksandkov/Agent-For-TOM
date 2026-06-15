import os
import docx
from docx.shared import Pt
from docx.shared import Cm as DocxCm  # Унікальна назва для відступів у Word
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.units import cm as PdfCm  # Унікальна назва для відступів у PDF

def create_docx(filename="Lab_Template_Final.docx"):
    """Генерує DOCX файл з ідеальним форматуванням (Times New Roman, 14pt, інтервал 1.5)"""
    doc = docx.Document()
    
    # Налаштування стилю "Normal" для всього документа
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    # Налаштування інтервалу та вирівнювання за замовчуванням
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.first_line_indent = DocxCm(1.25) # Червоний рядок
    paragraph_format.space_after = Pt(0)

    # 1. Заголовок "ЛАБОРАТОРНА РОБОТА" (по центру, звичайний)
    p_header = doc.add_paragraph()
    p_header.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header.paragraph_format.first_line_indent = DocxCm(0)
    p_header.add_run("ЛАБОРАТОРНА РОБОТА № [ВСТАВТЕ НОМЕР]")

    # 2. Назва роботи (по центру, жирний)
    p_title = doc.add_paragraph()
    p_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.first_line_indent = DocxCm(0)
    run_title = p_title.add_run("[Вставте назву лабораторної роботи]")
    run_title.bold = True
    
    doc.add_paragraph() # Порожній рядок

    # 3. Мета роботи (з абзацу, "Мета роботи" - жирним)
    p_meta = doc.add_paragraph()
    run_meta = p_meta.add_run("Мета роботи")
    run_meta.bold = True
    p_meta.add_run(" - [вставте текст мети роботи]")

    doc.add_paragraph()

    # 4. Загальні відомості (по центру, жирний)
    p_zagalni = doc.add_paragraph()
    p_zagalni.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_zagalni.paragraph_format.first_line_indent = DocxCm(0)
    p_zagalni.add_run("Загальні відомості").bold = True

    # Текст загальних відомостей (з абзацу, по ширині)
    doc.add_paragraph("[Вставте теоретичний матеріал тут. Абзацний відступ та вирівнювання по ширині налаштовані автоматично.]")
    
    doc.add_paragraph()

    # 5. Завдання. (по центру, жирний)
    p_zavd = doc.add_paragraph()
    p_zavd.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_zavd.paragraph_format.first_line_indent = DocxCm(0)
    p_zavd.add_run("Завдання.").bold = True

    # Пункти завдання (з абзацу)
    doc.add_paragraph("1. [Вставте перше завдання]")
    doc.add_paragraph("2. [Вставте друге завдання]")

    doc.add_paragraph()

    # 6. Контрольні запитання (по центру, жирний)
    p_kontr = doc.add_paragraph()
    p_kontr.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_kontr.paragraph_format.first_line_indent = DocxCm(0)
    p_kontr.add_run("Контрольні запитання").bold = True

    doc.add_paragraph("1. [Вставте перше запитання]")
    doc.add_paragraph("2. [Вставте друге запитання]")

    doc.add_paragraph()

    # 7. Література (по центру, жирний)
    p_lit = doc.add_paragraph()
    p_lit.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_lit.paragraph_format.first_line_indent = DocxCm(0)
    p_lit.add_run("Література").bold = True

    doc.add_paragraph("1. [Вставте перше джерело]")
    doc.add_paragraph("2. [Вставте друге джерело]")

    doc.save(filename)
    print(f"Успішно створено файл: {filename}")


def create_pdf(filename="Lab_Template_Final.pdf"):
    """Генерує PDF з точним дотриманням відступів та шрифтів"""
    
    # Реєстрація шрифтів Times New Roman (стандартні шляхи для Windows)
    try:
        pdfmetrics.registerFont(TTFont('TimesNewRoman', 'times.ttf'))
        pdfmetrics.registerFont(TTFont('TimesNewRoman-Bold', 'timesbd.ttf'))
        font_regular = 'TimesNewRoman'
        font_bold = 'TimesNewRoman-Bold'
    except Exception as e:
        print("Попередження: Шрифти Times New Roman не знайдено в системі. Використовується шрифт за замовчуванням.")
        font_regular = 'Helvetica'
        font_bold = 'Helvetica-Bold'

    # Використовуємо PdfCm для безпечного встановлення полів
    doc = SimpleDocTemplate(filename, pagesize=A4,
                            rightMargin=1.5*PdfCm, leftMargin=3.0*PdfCm,
                            topMargin=2.0*PdfCm, bottomMargin=2.0*PdfCm)
    
    # Налаштування стилів (14pt, 1.5 інтервал (leading ~21))
    style_center = ParagraphStyle(
        name='CenterNormal', fontName=font_regular, fontSize=14, leading=21, 
        alignment=TA_CENTER, spaceAfter=14)
    
    style_center_bold = ParagraphStyle(
        name='CenterBold', fontName=font_bold, fontSize=14, leading=21, 
        alignment=TA_CENTER, spaceAfter=14, spaceBefore=14)
    
    style_body = ParagraphStyle(
        name='BodyJustify', fontName=font_regular, fontSize=14, leading=21, 
        alignment=TA_JUSTIFY, firstLineIndent=35, spaceAfter=10) # 35 point ~ 1.25 cm

    story = []

    # 1. Заголовок
    story.append(Paragraph("ЛАБОРАТОРНА РОБОТА № [ВСТАВТЕ НОМЕР]", style_center))
    # 2. Назва
    story.append(Paragraph("[Вставте назву лабораторної роботи]", style_center_bold))
    
    # 3. Мета
    story.append(Paragraph("<b>Мета роботи</b> - [вставте текст мети роботи]", style_body))
    
    # 4. Загальні відомості
    story.append(Paragraph("Загальні відомості", style_center_bold))
    story.append(Paragraph("[Вставте теоретичний матеріал тут. Абзацний відступ та вирівнювання по ширині налаштовані автоматично.]", style_body))
    
    # 5. Завдання
    story.append(Paragraph("Завдання.", style_center_bold))
    story.append(Paragraph("1. [Вставте перше завдання]", style_body))
    story.append(Paragraph("2. [Вставте друге завдання]", style_body))
    
    # 6. Контрольні запитання
    story.append(Paragraph("Контрольні запитання", style_center_bold))
    story.append(Paragraph("1. [Вставте перше запитання]", style_body))
    story.append(Paragraph("2. [Вставте друге запитання]", style_body))

    # 7. Література
    story.append(Paragraph("Література", style_center_bold))
    story.append(Paragraph("1. [Вставте перше джерело]", style_body))
    story.append(Paragraph("2. [Вставте друге джерело]", style_body))

    doc.build(story)
    print(f"Успішно створено файл: {filename}")

if __name__ == "__main__":
    create_docx("Lab_Template_Final.docx")
    create_pdf("Lab_Template_Final.pdf")