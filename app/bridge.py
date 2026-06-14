"""
AppBridge — QObject that exposes Python slots and signals to QML.
Handles: session listing, template listing, navigation, mock generation.
"""
import json
import os
import shutil
import sys
import threading
import uuid
from datetime import datetime
from pathlib import Path

# pyrefly: ignore [missing-import]
from PyQt6.QtCore import QObject, pyqtSignal, pyqtSlot, pyqtProperty, QTimer, QUrl
# pyrefly: ignore [missing-import]
from PyQt6.QtWidgets import QFileDialog

# Make sure `app.backend.*` is importable when the bridge is loaded
# outside the normal `python main.py` launcher (e.g. by tests).
_REPO_ROOT = Path(__file__).resolve().parent.parent
if str(_REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(_REPO_ROOT))

from app.backend.pipeline import (  # noqa: E402
    TRANSIT_DIR,
    get_hf_token,
    load_env,
)
from app.backend.pipeline.bridge_adapter import (  # noqa: E402
    BridgePipelineAdapter,
)


# ─── Mock data ───────────────────────────────────────────────────────────────

MOCK_SESSIONS = [
    {
        "id": "s1", "name": "Лаба 1 — сортування масивів",
        "status": "completed", "template": "lab1", "hardness": "university_1",
        "duration": "14.3s", "created_at": "12 чер 2026, 14:23",
    },
    {
        "id": "s2", "name": "Звіт з практики — БД",
        "status": "completed", "template": "lab2", "hardness": "bachelor",
        "duration": "28.7s", "created_at": "11 чер 2026, 09:15",
    },
    {
        "id": "s3", "name": "Лаба 3 — графи, обхід у ширину",
        "status": "processing", "template": "lab1", "hardness": "university_1",
        "duration": "~10s", "created_at": "12 чер 2026, 14:28",
    },
    {
        "id": "s4", "name": "Тест — власний шаблон",
        "status": "warning", "template": "custom", "hardness": "school",
        "duration": "5.2s", "created_at": "10 чер 2026, 18:42",
    },
    {
        "id": "s5", "name": "Помилковий тест — невал. Python",
        "status": "failed", "template": "lab1", "hardness": "university_2",
        "duration": "3.1s", "created_at": "10 чер 2026, 18:30",
    },
    {
        "id": "s6", "name": "Чернетка — Лаба 2",
        "status": "draft", "template": "lab2", "hardness": "",
        "duration": "—", "created_at": "10 чер 2026, 15:11",
    },
]

MOCK_TEMPLATES = [
    {
        "id": "lab1", "name": "lab1", "display_name": "Лабораторна робота №1",
        "description": "Звіт про виконання лабораторної роботи з алгоритмів сортування",
        "is_builtin": True, "has_instructions": True,
    },
    {
        "id": "lab2", "name": "lab2", "display_name": "Лабораторна робота №2",
        "description": "Звіт з бази даних — проектування, нормалізація, SQL",
        "is_builtin": True, "has_instructions": True,
    },
    {
        "id": "custom1", "name": "custom1", "display_name": "Мій звіт",
        "description": "Користувацький шаблон без інструкцій",
        "is_builtin": False, "has_instructions": False,
    },
]

MOCK_INSTRUCTIONS = [
    {
        "id": "i1", "name": "Глобальні інструкції", "type": "global",
        "attached_to": None, "is_active": True,
        "created_at": "01 чер 2026",
    },
    {
        "id": "i2", "name": "Інструкції для Лаби 1", "type": "special",
        "attached_to": "Лабораторна робота №1", "is_active": True,
        "created_at": "02 чер 2026",
    },
    {
        "id": "i3", "name": "Інструкції для Лаби 2", "type": "special",
        "attached_to": "Лабораторна робота №2", "is_active": True,
        "created_at": "03 чер 2026",
    },
    {
        "id": "i4", "name": "Мій стиль написання", "type": "user_created",
        "attached_to": None, "is_active": True,
        "created_at": "05 чер 2026",
    },
]


class AppBridge(QObject):
    """
    Python ↔ QML bridge.
    Exposes mock data and pipeline simulation via Qt signals/slots.
    """

    # ─── Signals (Python → QML) ───────────────────────────────────────────────
    sessionsChanged = pyqtSignal()
    templatesChanged = pyqtSignal()
    instructionsChanged = pyqtSignal()
    currentScreenChanged = pyqtSignal()
    filesUpdated = pyqtSignal(str)
    fileWarning = pyqtSignal(str)   # warning message for the user
    isDarkThemeChanged = pyqtSignal()
    isSidebarCollapsedChanged = pyqtSignal()
    isBigFontChanged = pyqtSignal()
    sessionPayloadJsonChanged = pyqtSignal()

    # Pipeline signals
    pipelineStarted = pyqtSignal()
    pipelineProgress = pyqtSignal(int, str)        # percent, step_name
    pipelineLog = pyqtSignal(str, str)              # timestamp, message
    pipelineStepDone = pyqtSignal(int, str, str)   # step_index, name, detail
    pipelineStepActive = pyqtSignal(int)            # step_index
    pipelineFinished = pyqtSignal(str)              # session_id
    pipelineError = pyqtSignal(str, str)            # stage, message

    navigationRequest = pyqtSignal(str, str)        # screen, param

    def __init__(self, parent=None):
        super().__init__(parent)
        self._sessions = list(MOCK_SESSIONS)
        self._templates = list(MOCK_TEMPLATES)
        self._instructions = list(MOCK_INSTRUCTIONS)
        self._current_screen = "documents"
        self._pipeline_timer = None
        self._pipeline_step = 0
        self._uploaded_files = []
        self._prefs_file = os.path.join(os.path.dirname(__file__), "config", "user_preferences.json")
        self._transit_file = os.path.join(os.path.dirname(__file__), "transit", "userInput.json")
        self._is_dark_theme = False
        self._is_sidebar_collapsed = False
        self._is_big_font = False
        self._cache_ttl_days = 30
        self._session_retention_days = 90
        self._user_style_text = ""
        self._session_payload_json = "{}"
        self._db = None
        self._repo = None
        # Real pipeline state (replaces the previous mock timer).
        self._active_session_id: str | None = None
        self._pipeline_adapter: "BridgePipelineAdapter | None" = None
        self._pipeline_thread = None
        self._load_preferences()
        self._load_user_input()

    def _load_user_input(self):
        if os.path.exists(self._transit_file):
            try:
                with open(self._transit_file, "r", encoding="utf-8") as f:
                    self._session_payload_json = f.read()
                    if self._session_payload_json:
                        payload = json.loads(self._session_payload_json)
                        if "uploadedFiles" in payload and isinstance(payload["uploadedFiles"], list):
                            self._uploaded_files = payload["uploadedFiles"]
            except Exception as e:
                print(f"Error loading user input: {e}")

    @pyqtSlot()
    def clearTransitFolder(self):
        """Delete all files in transit folder."""
        transit_dir = os.path.join(os.path.dirname(__file__), "transit")
        if not os.path.exists(transit_dir):
            return
        for file in os.listdir(transit_dir):
            file_path = os.path.join(transit_dir, file)
            try:
                if os.path.isfile(file_path):
                    os.remove(file_path)
            except Exception as e:
                print(f"Error deleting file {file_path}: {e}")
        
        # Also clear the bridge's internal uploaded files list
        self._uploaded_files = []
        self.filesUpdated.emit("[]")

    @pyqtSlot(result=str)
    def getHfToken(self):
        """Read HF_TOKEN from .env file."""
        env_file = os.path.join(os.path.dirname(os.path.dirname(__file__)), ".env")
        if not os.path.exists(env_file):
            return ""
        try:
            with open(env_file, "r", encoding="utf-8") as f:
                for line in f:
                    if line.startswith("HF_TOKEN="):
                        return line.split("=", 1)[1].strip()
        except Exception as e:
            print(f"Error reading .env: {e}")
        return ""

    @pyqtSlot(result=str)
    def getVideoUrl(self):
        """Return the absolute file URL to the tutorial video."""
        video_path = os.path.join(os.path.dirname(__file__), "assets", "videos", "tutorial.mp4")
        return QUrl.fromLocalFile(video_path).toString()

    @pyqtSlot(str)
    def saveHfToken(self, token):
        """Save HF_TOKEN to .env file."""
        env_file = os.path.join(os.path.dirname(os.path.dirname(__file__)), ".env")
        lines = []
        token_found = False
        try:
            if os.path.exists(env_file):
                with open(env_file, "r", encoding="utf-8") as f:
                    lines = f.readlines()
            
            for i, line in enumerate(lines):
                if line.startswith("HF_TOKEN="):
                    lines[i] = f"HF_TOKEN={token}\n"
                    token_found = True
                    break
            
            if not token_found:
                # Ensure the file ends with a newline before appending
                if lines and not lines[-1].endswith("\n"):
                    lines[-1] += "\n"
                lines.append(f"HF_TOKEN={token}\n")
                
            with open(env_file, "w", encoding="utf-8") as f:
                f.writelines(lines)
        except Exception as e:
            print(f"Error saving to .env: {e}")

    def _save_user_input(self):
        try:
            os.makedirs(os.path.dirname(self._transit_file), exist_ok=True)
            if self._session_payload_json:
                payload = json.loads(self._session_payload_json)
                if "uploadedFiles" in payload and isinstance(payload["uploadedFiles"], list):
                    payload["uploadedFiles"] = [f for f in payload["uploadedFiles"] if f.get("status") == "done"]
                filtered_json = json.dumps(payload, ensure_ascii=False)
            else:
                filtered_json = "{}"
            with open(self._transit_file, "w", encoding="utf-8") as f:
                f.write(filtered_json)
        except Exception as e:
            print(f"Error saving user input: {e}")

    def _load_preferences(self):
        if os.path.exists(self._prefs_file):
            try:
                with open(self._prefs_file, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    self._is_dark_theme = data.get("isDarkTheme", False)
                    self._is_sidebar_collapsed = data.get("isSidebarCollapsed", False)
                    self._is_big_font = data.get("isBigFont", False)
                    self._cache_ttl_days = data.get("cacheTtlDays", 30)
                    self._session_retention_days = data.get("sessionRetentionDays", 90)
                    self._user_style_text = data.get("userStyleText", "")
            except Exception as e:
                print(f"Error loading preferences: {e}")

    def _save_preferences(self):
        try:
            os.makedirs(os.path.dirname(self._prefs_file), exist_ok=True)
            data = {
                "isDarkTheme": self._is_dark_theme,
                "isSidebarCollapsed": self._is_sidebar_collapsed,
                "isBigFont": self._is_big_font,
                "cacheTtlDays": self._cache_ttl_days,
                "sessionRetentionDays": self._session_retention_days,
                "userStyleText": self._user_style_text,
            }
            with open(self._prefs_file, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=4, ensure_ascii=False)
        except Exception as e:
            print(f"Error saving preferences: {e}")

    @pyqtSlot(str, int, int)
    def saveSettings(self, user_style: str, cache_ttl: int, session_retention: int):
        """Save all settings from the Settings screen at once."""
        self._user_style_text = user_style
        self._cache_ttl_days = max(1, cache_ttl)
        self._session_retention_days = max(1, session_retention)
        self._save_preferences()

    @pyqtSlot(result=str)
    def getUserStyle(self):
        return self._user_style_text

    @pyqtSlot(result=int)
    def getCacheTtlDays(self):
        return self._cache_ttl_days

    @pyqtSlot(result=int)
    def getSessionRetentionDays(self):
        return self._session_retention_days

    @pyqtSlot()
    def clearCache(self):
        """Manually clear all cached files (LLM, image, document caches)."""
        cache_dir = os.path.join(os.path.dirname(__file__), "cache")
        if os.path.exists(cache_dir):
            for item in os.listdir(cache_dir):
                item_path = os.path.join(cache_dir, item)
                try:
                    if os.path.isfile(item_path):
                        os.remove(item_path)
                    elif os.path.isdir(item_path):
                        shutil.rmtree(item_path)
                except Exception as e:
                    print(f"Error clearing cache item {item}: {e}")
        print("[bridge] Cache cleared manually")

    # ─── Properties ───────────────────────────────────────────────────────────

    @pyqtProperty(bool, notify=isDarkThemeChanged)
    def isDarkTheme(self):
        return self._is_dark_theme

    @isDarkTheme.setter
    def isDarkTheme(self, val):
        if self._is_dark_theme != val:
            self._is_dark_theme = val
            self.isDarkThemeChanged.emit()
            self._save_preferences()

    @pyqtProperty(bool, notify=isSidebarCollapsedChanged)
    def isSidebarCollapsed(self):
        return self._is_sidebar_collapsed

    @isSidebarCollapsed.setter
    def isSidebarCollapsed(self, val):
        if self._is_sidebar_collapsed != val:
            self._is_sidebar_collapsed = val
            self.isSidebarCollapsedChanged.emit()
            self._save_preferences()

    @pyqtProperty(bool, notify=isBigFontChanged)
    def isBigFont(self):
        return self._is_big_font

    @isBigFont.setter
    def isBigFont(self, val):
        if self._is_big_font != val:
            self._is_big_font = val
            self.isBigFontChanged.emit()
            self._save_preferences()

    @pyqtProperty(str, notify=sessionPayloadJsonChanged)
    def sessionPayloadJson(self):
        return self._session_payload_json

    @sessionPayloadJson.setter
    def sessionPayloadJson(self, val):
        if self._session_payload_json != val:
            self._session_payload_json = val
            self.sessionPayloadJsonChanged.emit()
            self._save_user_input()
            if val:
                try:
                    payload = json.loads(val)
                    if "uploadedFiles" in payload and isinstance(payload["uploadedFiles"], list):
                        self._uploaded_files = payload["uploadedFiles"]
                except Exception:
                    pass

    @pyqtProperty(str, notify=currentScreenChanged)
    def currentScreen(self):
        return self._current_screen

    # ─── Slots (QML → Python) ─────────────────────────────────────────────────

    @pyqtSlot(result=str)
    def getSessions(self):
        """Return all sessions as JSON string."""
        return json.dumps(self._sessions, ensure_ascii=False)

    @pyqtSlot(str, result=str)
    def getSessionsFiltered(self, status_filter: str):
        """Return sessions filtered by status ('all', 'completed', 'failed', 'draft')."""
        if status_filter == "all":
            data = self._sessions
        else:
            data = [s for s in self._sessions if s["status"] == status_filter]
        return json.dumps(data, ensure_ascii=False)

    @pyqtSlot(str, result=str)
    def getSessionById(self, session_id: str):
        for s in self._sessions:
            if s["id"] == session_id:
                return json.dumps(s, ensure_ascii=False)
        return "{}"

    @pyqtSlot(str)
    def deleteSession(self, session_id: str):
        self._sessions = [s for s in self._sessions if s["id"] != session_id]
        self.sessionsChanged.emit()

    @pyqtSlot(str)
    def duplicateSession(self, session_id: str):
        for s in self._sessions:
            if s["id"] == session_id:
                new_s = dict(s)
                new_s["id"] = str(uuid.uuid4())[:8]
                new_s["name"] = s["name"] + " — копія"
                new_s["status"] = "draft"
                new_s["duration"] = "—"
                new_s["created_at"] = datetime.now().strftime("%d %b %Y, %H:%M")
                self._sessions.insert(0, new_s)
                self.sessionsChanged.emit()
                return

    @pyqtSlot(str)
    def restoreSession(self, session_id: str):
        self._init_db()
        session = self._repo.sessions.get(session_id)
        if session:
            payload = json.loads(session.get("input_snapshot") or "{}")
            # Fill in required UI payload properties
            ui_payload = {
                "templateSlug": payload.get("template_slug", ""),
                "hardness": payload.get("hardness", ""),
                "imageMode": payload.get("image_mode", "none"),
                "documentLength": payload.get("document_length", "medium"),
                "useGlobalStyle": payload.get("use_global_style", True),
                "useGlobalInstructions": payload.get("use_global_instructions", True),
                "uploadedFiles": []
            }
            # Optional fields
            for k in ["theme", "target", "specialValue1", "specialValue2", "specialValue3"]:
                if k in payload:
                    ui_payload[k] = payload[k]
                    
            self.sessionPayloadJson = json.dumps(ui_payload, ensure_ascii=False)
            self.navigationRequest.emit("new_document", "")

    @pyqtSlot(result=str)
    def getTemplates(self):
        return json.dumps(self._templates, ensure_ascii=False)

    @pyqtSlot(result=str)
    def getLibraryFiles(self):
        self._init_db()
        files = self._repo.library_file.list_all()
        res = []
        for f in files:
            res.append({
                "id": f["id"],
                "original_name": f.get("original_name", "Unknown"),
                "file_size_bytes": f.get("file_size_bytes", 0),
                "created_at": f.get("created_at", "")[:10],
                "path": f.get("stored_path", "")
            })
        return json.dumps(res, ensure_ascii=False)

    @pyqtSlot(result=str)
    def getGeneratedFiles(self):
        self._init_db()
        sessions = self._repo.sessions.list_by_status("completed")
        res = []
        for s in sessions:
            res.append({
                "session_name": s.get("name", "Unknown Session"),
                "docx_path": s.get("docx_path", ""),
                "pdf_path": s.get("pdf_path", ""),
                "created_at": s.get("created_at", "")[:10]
            })
        return json.dumps(res, ensure_ascii=False)

    @pyqtSlot(str)
    def openFileExternal(self, rel_path: str):
        if not rel_path:
            return
        import subprocess
        full_path = os.path.abspath(os.path.join(os.path.dirname(os.path.dirname(__file__)), rel_path))
        if os.path.exists(full_path):
            if sys.platform == "win32":
                os.startfile(full_path)
            elif sys.platform == "darwin":
                subprocess.call(["open", full_path])
            else:
                subprocess.call(["xdg-open", full_path])

    @pyqtSlot(str)
    def showInFolder(self, rel_path: str):
        if not rel_path:
            return
        import subprocess
        full_path = os.path.abspath(os.path.join(os.path.dirname(os.path.dirname(__file__)), rel_path))
        if os.path.exists(full_path):
            if sys.platform == "win32":
                subprocess.call(['explorer', '/select,', full_path])
            elif sys.platform == "darwin":
                subprocess.call(["open", "-R", full_path])
            else:
                subprocess.call(["xdg-open", os.path.dirname(full_path)])

    def _init_db(self):
        if self._db is None:
            from app.backend.db.connection import Database
            from app.backend.db.facade import BridgeRepository
            self._db = Database()
            self._repo = BridgeRepository(self._db)

    def _fetch_instructions(self) -> list[dict]:
        self._init_db()
        
        # Get the latest version of each instruction (type, template_id combination)
        # so that we can show them even if is_active = 0 (toggled off)
        query = """
            SELECT * FROM instructions 
            WHERE id IN (
                SELECT id FROM instructions 
                GROUP BY type, IFNULL(template_id, 'global')
                HAVING version = MAX(version)
            )
        """
        rows = self._db.conn.execute(query).fetchall()
        items = [dict(r) for r in rows]
        
        templates = {t["id"]: t for t in self._repo.templates.list_all()}
        
        active_categories = set()
        for i in items:
            if i["is_active"]:
                active_categories.add((i["type"], i.get("template_id")))
                
        res = []
        for i in items:
            t_id = i.get("template_id")
            t_name = templates.get(t_id, {}).get("display_name") if t_id else None
            
            name = ""
            if i["type"] == "global":
                name = "Глобальні інструкції"
            elif i["type"] == "special":
                name = f"Інструкції для {t_name}" if t_name else "Спеціальні інструкції"
            else:
                name = "Мій стиль написання"

            cat = i.get("created_at", "")[:10]
            is_act = bool(i["is_active"])
            
            can_pin = True
            if not is_act and (i["type"], t_id) in active_categories:
                can_pin = False
            
            res.append({
                "id": i["id"],
                "name": name,
                "type": i["type"],
                "attached_to": t_name,
                "is_active": is_act,
                "can_pin": can_pin,
                "created_at": cat,
                "content": i.get("content", "")
            })
            
        def _sort_key(x):
            is_global = 0 if x["type"] == "global" else 1
            is_pinned = 0 if x["is_active"] else 1
            type_prio = {"global": 0, "special": 1, "user_created": 2}.get(x["type"], 99)
            return (is_global, is_pinned, type_prio, x["name"])
            
        res.sort(key=_sort_key)
        
        return res

    @pyqtSlot(result=str)
    def getInstructions(self):
        return json.dumps(self._fetch_instructions(), ensure_ascii=False)

    @pyqtSlot(str, result=str)
    def getInstructionsFiltered(self, type_filter: str):
        instructions = self._fetch_instructions()
        if type_filter == "all":
            data = instructions
        elif type_filter == "global":
            data = [i for i in instructions if i["type"] == "global"]
        elif type_filter == "special":
            data = [i for i in instructions if i["type"] == "special"]
        elif type_filter == "user_created":
            data = [i for i in instructions if i["type"] == "user_created"]
        elif type_filter == "unattached":
            data = [i for i in instructions if i.get("attached_to") is None]
        else:
            data = instructions
        return json.dumps(data, ensure_ascii=False)

    @pyqtSlot(str, bool)
    def toggleInstructionStatus(self, inst_id: str, is_active: bool):
        self._init_db()
        if not is_active:
            self._repo.instructions.deactivate(inst_id)
        else:
            inst = self._repo.instructions.get(inst_id)
            if inst:
                template_id = inst["template_id"]
                type_ = inst["type"]
                with self._db.transaction():
                    if template_id is None:
                        self._db.conn.execute("UPDATE instructions SET is_active = 0 WHERE type = ? AND template_id IS NULL AND is_active = 1", (type_,))
                    else:
                        self._db.conn.execute("UPDATE instructions SET is_active = 0 WHERE type = ? AND template_id = ? AND is_active = 1", (type_, template_id))
                    self._db.conn.execute("UPDATE instructions SET is_active = 1 WHERE id = ?", (inst_id,))
        self.instructionsChanged.emit()

    @pyqtSlot(str, str, str)
    def createInstruction(self, type_: str, template_id: str, content: str):
        self._init_db()
        self._repo.instructions.save_new_version(
            type_=type_,
            content=content,
            template_id=template_id if template_id else None
        )
        self.instructionsChanged.emit()

    @pyqtSlot(str)
    def deleteInstruction(self, inst_id: str):
        self._init_db()
        inst = self._repo.instructions.get(inst_id)
        if inst:
            type_ = inst["type"]
            template_id = inst["template_id"]
            if type_ == "global":
                return # Cannot delete global
            with self._db.transaction():
                if template_id is None:
                    self._db.conn.execute("DELETE FROM instructions WHERE type = ? AND template_id IS NULL", (type_,))
                else:
                    self._db.conn.execute("DELETE FROM instructions WHERE type = ? AND template_id = ?", (type_, template_id))
        self.instructionsChanged.emit()

    @pyqtSlot(str)
    def navigate(self, screen: str):
        """Navigate to a screen by name."""
        self._current_screen = screen
        self.currentScreenChanged.emit()
        self.navigationRequest.emit(screen, "")

    @pyqtSlot(str, str)
    def navigateWithParam(self, screen: str, param: str):
        self._current_screen = screen
        self.currentScreenChanged.emit()
        self.navigationRequest.emit(screen, param)

    @pyqtSlot()
    def openFileDialog(self):
        # pyrefly: ignore [missing-import]
        from PyQt6.QtCore import QStandardPaths
        default_dir = QStandardPaths.writableLocation(QStandardPaths.StandardLocation.DocumentsLocation)
        if not default_dir:
            default_dir = QStandardPaths.writableLocation(QStandardPaths.StandardLocation.HomeLocation)
            
        files, _ = QFileDialog.getOpenFileNames(
            None, 
            "Оберіть файли", 
            default_dir, 
            "Supported Files (*.pdf *.docx *.pptx *.png *.jpg *.jpeg)"
        )
        if files:
            urls = [QUrl.fromLocalFile(f).toString() for f in files]
            self.uploadFilesUrls(urls)

    @pyqtSlot('QVariantList')
    def uploadFilesUrls(self, urls):
        import threading
        thread = threading.Thread(target=self._process_files_sequentially, args=(urls,), daemon=True)
        thread.start()

    def _process_files_sequentially(self, urls):
        import time
        time.sleep(1)  # Delay after dialog closes
        
        app_dir = os.path.dirname(__file__)
        transit_dir = os.path.join(app_dir, "transit")
        os.makedirs(transit_dir, exist_ok=True)

        loaded_names = {f["name"] for f in self._uploaded_files}
        duplicates = []
        invalid_types = []

        for url in urls:
            if isinstance(url, QUrl):
                local_path = url.toLocalFile()
            else:
                local_path = QUrl(str(url)).toLocalFile()

            if not local_path or not os.path.exists(local_path):
                continue

            filename = os.path.basename(local_path)
            ext = os.path.splitext(filename)[1].lower()

            # Duplicate check
            if filename in loaded_names:
                duplicates.append(filename)
                continue

            # Type check
            if ext not in ['.pdf', '.docx', '.pptx', '.png', '.jpg', '.jpeg']:
                invalid_types.append(filename)
                continue

            try:
                size_bytes = os.path.getsize(local_path)
            except Exception as e:
                print(f"Error getting size for {filename}: {e}")
                continue

            # Human-readable size
            if size_bytes < 1024:
                meta = f"{size_bytes} Б"
            elif size_bytes < 1024 * 1024:
                meta = f"{size_bytes / 1024:.1f} КБ"
            else:
                meta = f"{size_bytes / (1024 * 1024):.1f} МБ"

            import random
            import string
            import subprocess

            rand_chars = ''.join(random.choices(string.ascii_letters + string.digits, k=5))
            txt_filename = f"{filename}_{rand_chars}.txt"
            attached_dir = os.path.join(transit_dir, "attached")
            os.makedirs(attached_dir, exist_ok=True)
            txt_path = os.path.join(attached_dir, txt_filename)

            # Add to UI as processing
            file_record = {
                "name": filename,
                "meta": meta,
                "path": local_path,
                "txt_filename": txt_filename,
                "symbols": 0,
                "status": "processing"
            }
            self._uploaded_files.append(file_record)
            self.filesUpdated.emit(json.dumps(self._uploaded_files, ensure_ascii=False))

            def update_status(status_str, sym_count=0):
                for f in self._uploaded_files:
                    if f["name"] == filename:
                        f["status"] = status_str
                        if sym_count > 0:
                            f["symbols"] = sym_count
                        break

            # 50MB limit check
            if size_bytes > 50 * 1024 * 1024:
                update_status("error")
                self.filesUpdated.emit(json.dumps(self._uploaded_files, ensure_ascii=False))
                self.fileWarning.emit(f"Файл перевищує ліміт у 50 МБ: {filename}")
                continue

            # Convert using backend scripts
            converter_script = None
            if ext == '.pdf': converter_script = 'pdf2txt.py'
            elif ext == '.docx': converter_script = 'docx2txt.py'
            elif ext == '.pptx': converter_script = 'pptx2txt.py'
            elif ext in ['.png', '.jpg', '.jpeg']: converter_script = 'image2txt.py'
            
            symbols = 0
            success = False
            
            if converter_script:
                script_path = os.path.join(app_dir, "backend", converter_script)
                try:
                    subprocess.run(
                        [sys.executable, script_path, "-o", txt_path, local_path],
                        check=True,
                        capture_output=True,
                        text=True
                    )
                    success = True
                except subprocess.CalledProcessError as e:
                    print(f"Error converting {filename}: {e.stderr}")
                    update_status("error")
                    self.filesUpdated.emit(json.dumps(self._uploaded_files, ensure_ascii=False))
                    self.fileWarning.emit(f"Помилка конвертації файлу: {filename}")
                    continue
            else:
                try:
                    with open(local_path, "r", encoding="utf-8", errors="replace") as f_in, \
                         open(txt_path, "w", encoding="utf-8") as f_out:
                         f_out.write(f_in.read())
                    success = True
                except Exception as e:
                    print(f"Error reading text file {filename}: {e}")
                    update_status("error")
                    self.filesUpdated.emit(json.dumps(self._uploaded_files, ensure_ascii=False))
                    self.fileWarning.emit(f"Помилка читання файлу: {filename}")
                    continue

            if success and os.path.exists(txt_path):
                try:
                    with open(txt_path, "r", encoding="utf-8") as f:
                        symbols = len(f.read())
                except Exception:
                    symbols = 0
                
                update_status("done", symbols)
                self.filesUpdated.emit(json.dumps(self._uploaded_files, ensure_ascii=False))
                continue

            loaded_names.add(filename)
            self.filesUpdated.emit(json.dumps(self._uploaded_files, ensure_ascii=False))

        warnings = []
        if duplicates:
            warnings.append(f"Файл(и) вже завантажені: {', '.join(duplicates)}")
        if invalid_types:
            warnings.append(f"Непідтримуваний формат: {', '.join(invalid_types)}")
        
        if warnings:
            self.fileWarning.emit(" | ".join(warnings))

        self.filesUpdated.emit(json.dumps(self._uploaded_files, ensure_ascii=False))

    @pyqtSlot(result=str)
    def getSessionResult(self):
        self._init_db()
        if not self._active_session_id:
            return self.getMockResult()
        session = self._repo.sessions.get(self._active_session_id)
        if not session:
            return self.getMockResult()

        token_usage = json.loads(session.get("token_usage") or "{}")
        duration_ms = session.get("duration_ms", 0)
        snapshot = json.loads(session.get("input_snapshot") or "{}")
        val = json.loads(session.get("validation_result") or "{}")
        
        data = {
            "session_name": session.get("name"),
            "status": session.get("status"),
            "error_stage": session.get("error_stage"),
            "error_message": session.get("error_message"),
            "duration": f"{duration_ms} ms" if duration_ms else "",
            "image_count": session.get("image_count", 0),
            "word_count": val.get("word_count", 0),
            "word_count_min": val.get("word_count_target_min", 0),
            "word_count_max": val.get("word_count_target_max", 0),
            "input_tokens": str(token_usage.get("text_model", {}).get("input_tokens", 0)),
            "output_tokens": str(token_usage.get("text_model", {}).get("output_tokens", 0)),
            "cached_tokens": str(token_usage.get("text_model", {}).get("cached_tokens", 0)),
            "template": snapshot.get("template_slug", ""),
            "hardness": snapshot.get("hardness", ""),
            "image_mode": snapshot.get("image_mode", ""),
            "warnings": val.get("warnings", []),
            "docx_path": session.get("docx_output", ""),
            "pdf_path": session.get("pdf_output", ""),
        }
        return json.dumps(data, ensure_ascii=False)

    @pyqtSlot(result=str)
    def getSessionFilledPy(self):
        self._init_db()
        if not self._active_session_id:
            return self.getMockFilledPy()
        session = self._repo.sessions.get(self._active_session_id)
        if not session:
            return self.getMockFilledPy()
        path = session.get("filled_py_path")
        if not path:
            return ""
        full_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), path)
        try:
            with open(full_path, "r", encoding="utf-8") as f:
                return f.read()
        except:
            return "Помилка читання файлу: " + full_path

    @pyqtSlot(str)
    def downloadFile(self, rel_path: str):
        if not rel_path:
            return
        full_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), rel_path)
        if not os.path.exists(full_path):
            print("File not found:", full_path)
            return
        
        file_name = os.path.basename(full_path)
        
        # Open save dialog
        default_dir = os.path.expanduser("~\\Downloads")
        save_path, _ = QFileDialog.getSaveFileName(
            None, "Зберегти файл", 
            os.path.join(default_dir, file_name)
        )
        if save_path:
            try:
                shutil.copy2(full_path, save_path)
            except Exception as e:
                print(f"Error copying file {full_path} to {save_path}: {e}")

    @pyqtSlot(str)
    def deleteFile(self, filename: str):
        """Remove a file from the uploaded list and delete from transit."""
        self._uploaded_files = [f for f in self._uploaded_files if f["name"] != filename]
        app_dir = os.path.dirname(__file__)
        transit_path = os.path.join(app_dir, "transit", filename)
        if os.path.exists(transit_path):
            try:
                os.remove(transit_path)
            except Exception as e:
                print(f"Could not delete transit file {filename}: {e}")
        self.filesUpdated.emit(json.dumps(self._uploaded_files, ensure_ascii=False))

    # ─── Pipeline simulation ──────────────────────────────────────────────────
    #
    # NOTE: The previous block (mock _PIPELINE_STEPS, _LOG_MESSAGES,
    # _run_step, _complete_step) has been replaced by a real
    # implementation that delegates to `app.backend.pipeline.orchestrator`.
    # The 6-step QML contract is preserved by emitting synthetic
    # `pipelineStepDone` signals for steps 0..5 inside BridgePipelineAdapter.

    # ─── Pipeline helpers ─────────────────────────────────────────────────────

    def _build_pipeline_payload(
        self,
        session_name: str,
        template_id: str,
        length: str,
        hardness: str,
        image_mode: str,
        goal: str,
    ) -> dict:
        """Merge the per-generation args with the cached `sessionPayloadJson`.

        Returns a single dict ready to be written as
        ``session_context.json`` inside a transit snapshot.
        """
        # Pull the cached UI payload (set via the `sessionPayloadJson` property).
        cached: dict = {}
        try:
            if self._session_payload_json:
                cached = json.loads(self._session_payload_json)
        except (ValueError, TypeError):
            cached = {}

        # Map QML hardness values to internal canonical values.
        hardness_map = {
            "school": "school",
            "bachelor": "bachelor",
            "university_1": "university_1",
            "university_2": "university_2",
            "master": "master",
        }
        canonical_hardness = hardness_map.get(hardness, hardness or "university_1")

        length_map = {
            "short": "short",
            "middle": "middle",
            "long": "long",
        }
        canonical_length = length_map.get(length, length or "middle")

        image_mode_map = {
            "none": "none",
            "references": "references",
            "full": "full",
        }
        canonical_image_mode = image_mode_map.get(image_mode, image_mode or "none")

        return {
            "template_id": template_id or cached.get("template_id") or "lab1",
            "name": session_name or cached.get("name") or "Нова сесія",
            "theme": cached.get("theme") or session_name or "",
            "goal": goal or cached.get("goal") or "",
            "length": canonical_length,
            "hardness": canonical_hardness,
            "image_mode": canonical_image_mode,
            "include_special_instructions": bool(
                cached.get("include_special_instructions", True)
            ),
            "include_user_style": bool(cached.get("include_user_style", False)),
            "user_input": cached.get("user_input") or "",
            "uploaded_files": cached.get("uploadedFiles") or cached.get("uploaded_files") or [],
            "gap_values": cached.get("gap_values") or {},
        }

    def _materialize_transit_snapshot(
        self,
        session_id: str,
        payload: dict,
    ) -> Path:
        """Write a complete transit snapshot for the orchestrator to read.

        Layout (under ``app/debug/transit/<session_id>/``):
          - session_context.json
          - general_instructions.md   (copied from app/instructions/)
          - <template>_fill.md        (copied from app/instructions/template-ins/)
          - <template>_params.json    (built from payload.gap_values)
          - library_files.json        (one entry per uploaded file)
          - context.json              (one entry per uploaded file)
          - attached/<hash>.txt       (copied/moved from app/transit/)
        """
        # Resolve directories.
        snap_dir = TRANSIT_DIR / session_id
        attached_dir = snap_dir / "attached"
        app_instructions_dir = Path(__file__).resolve().parent / "instructions"
        app_template_ins_dir = app_instructions_dir / "template-ins"
        app_transit_dir = Path(__file__).resolve().parent / "transit"
        # Use a stable, relative path the rest of the app understands.
        try:
            from app.backend.pipeline.utils import REPO_ROOT  # noqa: WPS433
        except Exception:
            REPO_ROOT = Path(__file__).resolve().parent.parent

        template_id = payload.get("template_id") or "lab1"
        attached_dir.mkdir(parents=True, exist_ok=True)

        # 1) session_context.json
        snapshot = dict(payload)
        snapshot["created_at"] = datetime.utcnow().isoformat() + "+00:00"
        snapshot["id"] = session_id
        snapshot["status"] = "processing"
        file_refs = {
            "global_instructions": "general_instructions.md",
            "template_instructions": f"{template_id}_fill.md",
            "template_params": f"{template_id}_params.json",
            "context": "context.json",
            "library_files": "library_files.json",
            "attached": "attached/",
        }
        # Build the gap_values_ref so stage1 picks up our params.
        snapshot["gap_values_ref"] = f"{template_id}_params.json"
        snapshot["global_instructions_hash"] = ""
        snapshot["style_hash"] = ""
        snapshot["attached_files"] = []  # populated below
        session_ctx = {
            "id": session_id,
            "name": payload.get("name") or "Нова сесія",
            "template_id": template_id,
            "status": "processing",
            "input_snapshot": snapshot,
            "file_refs": file_refs,
            "created_at": snapshot["created_at"],
        }
        (snap_dir / "session_context.json").write_text(
            json.dumps(session_ctx, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

        # 2) general_instructions.md
        src_g = app_instructions_dir / "global_instructions.md"
        if src_g.is_file():
            shutil.copy2(src_g, snap_dir / "general_instructions.md")
            session_ctx["input_snapshot"]["global_instructions_hash"] = self._hash_file(src_g)

        # 3) template instructions
        src_ti = app_template_ins_dir / f"{template_id}_fill.md"
        if not src_ti.is_file():
            # Fall back to lab1_fill.md if the requested template file is missing.
            src_ti = app_template_ins_dir / "lab1_fill.md"
        if src_ti.is_file():
            shutil.copy2(src_ti, snap_dir / f"{template_id}_fill.md")

        # 4) template params (gap_values, schema-aware)
        gap_values = payload.get("gap_values") or self._fallback_gap_values(payload)
        session_ctx["input_snapshot"]["gap_values"] = gap_values
        params_doc = {"gap_values": gap_values}
        (snap_dir / f"{template_id}_params.json").write_text(
            json.dumps(params_doc, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

        # 5) user_style.md (optional)
        if payload.get("include_user_style"):
            us = app_instructions_dir / "user_style.md"
            if us.is_file() and us.read_text(encoding="utf-8").strip():
                shutil.copy2(us, snap_dir / "user_style.md")
                session_ctx["input_snapshot"]["style_hash"] = self._hash_file(us)

        # 6) uploaded files: move/copy from app/transit/ into attached/
        uploaded = payload.get("uploaded_files") or []
        library_files: list[dict] = []
        context_files: list[dict] = []
        attached_hashes: list[str] = []
        for f_info in uploaded:
            if not isinstance(f_info, dict):
                continue
            if f_info.get("status") not in (None, "done"):
                continue
            fname = f_info.get("name")
            txt_filename = f_info.get("txt_filename")
            if not fname:
                continue
                
            if txt_filename:
                src = app_transit_dir / "attached" / txt_filename
            else:
                src = app_transit_dir / fname
                
            if not src.is_file():
                # Try with the file_hash as the stored name.
                fh = f_info.get("file_hash") or fname
                src = app_transit_dir / fh
            if not src.is_file():
                print(f"[bridge] warning: could not find source file for {fname} at {src}")
                continue
            file_hash = f_info.get("file_hash") or self._hash_file(src)
            dest = attached_dir / f"{file_hash}.txt"
            try:
                if dest.resolve() != src.resolve():
                    shutil.copy2(src, dest)
            except OSError as exc:
                print(f"[bridge] could not copy attached file {fname}: {exc}")
                continue
            attached_hashes.append(file_hash)
            original_type = f_info.get("type") or "text/plain"
            library_files.append({
                "id": f_info.get("id") or str(uuid.uuid4()),
                "original_name": fname,
                "original_type": original_type,
                "file_hash": file_hash,
                "original_sha256": file_hash,
                "stored_path": str(
                    (REPO_ROOT / "storage" / "library" / file_hash[:2] / f"{file_hash}.txt")
                ).replace("\\", "/"),
                "converted_text_path": f"attached/{file_hash}.txt",
                "conversion_status": "done",
                "file_size_bytes": src.stat().st_size if src.exists() else 0,
                "converted_at": datetime.utcnow().isoformat() + "+00:00",
                "created_at": datetime.utcnow().isoformat() + "+00:00",
                "last_used_at": datetime.utcnow().isoformat() + "+00:00",
            })
            context_files.append({
                "file_hash": file_hash,
                "original_name": fname,
                "original_type": original_type,
                "original_sha256": file_hash,
                "converted_text_path": f"attached/{file_hash}.txt",
                "token_count": 0,
                "was_summarized": False,
            })
        # Update snapshot with attached_files list now that we have it.
        session_ctx["input_snapshot"]["attached_files"] = attached_hashes
        (snap_dir / "session_context.json").write_text(
            json.dumps(session_ctx, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

        # 7) library_files.json + context.json
        (snap_dir / "library_files.json").write_text(
            json.dumps({"files": library_files}, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        (snap_dir / "context.json").write_text(
            json.dumps(
                {
                    "files": context_files,
                    "merged_text_preview": "",
                    "total_tokens": 0,
                    "cache_key": "",
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        return snap_dir

    @staticmethod
    def _hash_file(path: Path) -> str:
        import hashlib
        h = hashlib.sha256()
        try:
            with open(path, "rb") as fh:
                for chunk in iter(lambda: fh.read(65536), b""):
                    h.update(chunk)
        except OSError:
            pass
        return h.hexdigest()

    @staticmethod
    def _fallback_gap_values(payload: dict) -> dict:
        """Build a schema-aware `gap_values` block when the UI didn't provide one.

        Mirrors the keys in ``app/instructions/template-ins/lab1_fill.md``
        so that the gap_assembler can produce a real filled.py.
        """
        name = payload.get("name") or "Нова лабораторна робота"
        theme = payload.get("theme") or name
        goal = payload.get("goal") or "дослідити та проаналізувати поставлену задачу."
        user_input = payload.get("user_input") or ""
        return {
            "lab_number": {"value": "1", "ai_accessible": False},
            "work_title": {"value": theme, "ai_accessible": True},
            "goal": {"value": goal, "ai_accessible": True},
            "general_info": {
                "value": user_input or f"{theme} — базові теоретичні відомості.",
                "ai_accessible": True,
            },
            "tasks": {
                "value": [
                    "Реалізувати алгоритми відповідно до мети роботи.",
                    "Продемонструвати їх роботу на тестових даних.",
                    "Зробити висновки щодо отриманих результатів.",
                ],
                "ai_accessible": True,
            },
            "control_questions": {
                "value": [
                    "У чому полягає мета роботи?",
                    "Які основні кроки виконаних алгоритмів?",
                ],
                "ai_accessible": True,
            },
            "bibliography": {
                "value": [
                    "Кнут, Д. Е. Мистецтво програмування. Т. 3 : Сортування і пошук. Київ : Вільямс, 2020. 824 с.",
                ],
                "ai_accessible": True,
            },
        }

    def _update_session_from_run(self, run_dict: dict) -> None:
        """Push the orchestrator's result back into ``self._sessions``."""
        if not self._active_session_id:
            return
        index = run_dict.get("index") or {}
        status = index.get("status") or "completed"
        # Map orchestrator status to UI status vocabulary.
        ui_status = {
            "completed": "completed",
            "failed": "failed",
            "cancelled": "cancelled",
        }.get(status, "completed")
        duration_ms = index.get("duration_ms") or 0
        duration_s = max(0.0, duration_ms / 1000.0)
        for s in self._sessions:
            if s["id"] == self._active_session_id:
                s["status"] = ui_status
                s["duration"] = f"{duration_s:.1f}s" if duration_s else "—"
                s["docx_path"] = index.get("artifacts", {}).get("docx")
                s["pdf_path"] = index.get("artifacts", {}).get("pdf")
                break
        self.sessionsChanged.emit()

    @pyqtSlot(str)
    def startGeneration(self, payload_str: str):
        """Start a REAL pipeline generation.
        """
        import json
        payload = json.loads(payload_str)
        
        session_name = payload.get("documentName", "")
        template_id = payload.get("template_id", "lab1")
        length = payload.get("lengthMode", "middle")
        hardness = "university_1"
        image_mode = payload.get("image_mode", "none")
        goal = payload.get("documentGoal", "")
        include_user_style = payload.get("userStyleId") not in [None, "", "none"]

        self._init_db()

        # 1. New session in processing state.
        new_id = str(uuid.uuid4())[:8]
        new_session = {
            "id": new_id,
            "name": session_name or "Нова сесія",
            "status": "processing",
            "template": template_id,
            "hardness": hardness,
            "duration": "—",
            "created_at": datetime.now().strftime("%d %b %Y, %H:%M"),
        }
        self._sessions.insert(0, new_session)
        self.sessionsChanged.emit()
        self._active_session_id = new_id

        # 2. Payload.
        pipeline_payload = self._build_pipeline_payload(
            session_name, template_id, length, hardness, image_mode, goal
        )
        pipeline_payload["uploaded_files"] = payload.get("uploadedFiles", [])
        pipeline_payload["user_input"] = payload.get("sessionHints", "")
        pipeline_payload["include_user_style"] = include_user_style

        # 3. Snapshot.
        try:
            snap_dir = self._materialize_transit_snapshot(new_id, pipeline_payload)
        except Exception as exc:  # noqa: BLE001
            self.pipelineError.emit("stage1", f"snapshot failed: {exc}")
            for s in self._sessions:
                if s["id"] == new_id:
                    s["status"] = "failed"
                    break
            self.sessionsChanged.emit()
            return
            
        # DB insertion
        try:
            from app.backend.db.repositories.sessions import SessionCreate
            session_create = SessionCreate(
                id=new_id,
                name=session_name or "Нова сесія",
                status="processing",
                template_id=template_id,
                hardness=hardness,
                input_snapshot=json.dumps(pipeline_payload, ensure_ascii=False)
            )
            self._repo.sessions.create(session_create)
        except Exception as e:
            print(f"[bridge] Warning: DB insertion failed: {e}")

        # 4. Background thread + adapter.
        self._pipeline_adapter = BridgePipelineAdapter(self)
        self._pipeline_thread = threading.Thread(
            target=self._run_pipeline_blocking,
            args=(self._pipeline_adapter, snap_dir, new_id),
            daemon=True,
        )
        self._pipeline_thread.start()

    def _run_pipeline_blocking(
        self,
        adapter: "BridgePipelineAdapter",
        snap_dir: Path,
        session_id: str,
    ) -> None:
        """Thread entry point: run adapter, then push result into _sessions."""
        try:
            result = adapter.run(snap_dir)
            self._update_session_from_run(result)
        except Exception as exc:  # noqa: BLE001
            self.pipelineError.emit("orchestrator", f"{type(exc).__name__}: {exc}")
            for s in self._sessions:
                if s["id"] == session_id:
                    s["status"] = "failed"
                    break
            self.sessionsChanged.emit()

    @pyqtSlot()
    def cancelPipeline(self):
        """Request cancellation of the running pipeline (if any)."""
        if self._pipeline_adapter is not None:
            self._pipeline_adapter.request_cancel()
        if self._active_session_id:
            for s in self._sessions:
                if s["id"] == self._active_session_id:
                    s["status"] = "cancelled"
                    break
            self.sessionsChanged.emit()

    @pyqtSlot(result=str)
    def getMockResult(self):
        """Return mock result data for the result screen."""
        return json.dumps({
            "session_name": "Лаба 1 — сортування масивів",
            "status": "completed",
            "token_usage": {
                "text_model": {"input_tokens": 4821, "output_tokens": 2103, "cached_tokens": 3200},
                "image_model": {"images_generated": 2}
            },
            "duration_ms": 16400,
            "image_count": 2,
            "word_count": 1923,
            "word_count_min": 1700,
            "word_count_max": 2500,
            "warnings": ["Word count 3% нижче від мети"],
        }, ensure_ascii=False)

    @pyqtSlot(result=str)
    def getMockFilledPy(self):
        """Return mock filled.py content for syntax-highlighted preview."""
        return '''# custom_template_lab1.py
# Template: Лабораторна робота №1
# Auto-generated by Agent-for-TOM

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate

def create_docx(filename):
    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Міністерство освіти і науки України")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    doc.add_heading("Лабораторна робота №1", level=1)
    doc.add_paragraph("""
Мета роботи: порівняти ефективність бульбашкового сортування
та quicksort на масивах різного розміру. Провести часові
характеристики для масивів розміром N = 100, 1000, 10000.
""")

    [[IMAGE|type:diagram|subject:Порівняння часу сортування|context:Графік O(n²) vs O(n log n)|style:labeled]]

    doc.add_heading("Висновки", level=2)
    doc.add_paragraph("""
Quicksort демонструє значно кращу продуктивність на великих
масивах завдяки середній складності O(n log n) порівняно з
O(n²) у bubble sort.
""")
    doc.save(filename)

def create_pdf(filename):
    doc = SimpleDocTemplate(filename, pagesize=A4)
    # ... (аналогічний вміст)
    doc.build([])
'''
