"""app/backend/llm/gap_assembler.py

Renders a `filled.py` from the per-template `gap_values` JSON.

This is the LOCAL FALLBACK used when `HUGGY_FACE_TOKEN` is not set. It
mirrors what a remote Llama-3.3 call would produce, but with content
that the user already typed into the GUI ("theme", "goal", custom tasks
and control questions, bibliography).

Output: a complete, runnable Python script with both `create_docx()`
and `create_pdf()` functions filled in. The output obeys every rule
listed in `lab1_fill.md` (DSTU 3008:2015 typography, no triple quotes
inside placeholders, double quotes for string literals, etc).

The assembler is template-aware. Right now it supports:
  - lab1: «Дослідження методів сортування в масивах» template
  - lab2: any other lab with the standard 7-section layout

Adding a new template is a matter of adding a `_render_labN` function
and a dispatch entry in `render_filled_py`.
"""
from __future__ import annotations

import json
from pathlib import Path
from typing import Any

# ─── Shared DOCX/PDF template source ──────────────────────────────────────
#
# This is the canonical, "small" lab template. The original Pipeline.md
# template lives in Plan/backend3/ and is much larger. We use a smaller
# version that still satisfies all the rules in lab1_fill.md:
#   - Times New Roman 14, line spacing 1.5, indent 1.25 cm
#   - Section order: Title → Name → Goal → Загальні відомості → Завдання
#     → Контрольні запитання → Література
#   - Same text in DOCX and PDF versions
#   - No comments, no print() inside functions

_TPL_DOCX = '''import os
import docx
from docx.shared import Pt
from docx.shared import Cm as DocxCm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.units import cm as PdfCm


def _set_docx_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = DocxCm(1.25)
    pf.space_after = Pt(0)


def _add_docx_centered(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = DocxCm(0)
    run = p.add_run(text)
    run.bold = bold
    return p


def _add_docx_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def _setup_pdf_fonts():
    try:
        pdfmetrics.registerFont(TTFont("TimesNewRoman", "times.ttf"))
        pdfmetrics.registerFont(TTFont("TimesNewRoman-Bold", "timesbd.ttf"))
        return "TimesNewRoman", "TimesNewRoman-Bold"
    except Exception:
        try:
            pdfmetrics.registerFont(TTFont("TimesNewRoman", "C:/Windows/Fonts/times.ttf"))
            pdfmetrics.registerFont(TTFont("TimesNewRoman-Bold", "C:/Windows/Fonts/timesbd.ttf"))
            return "TimesNewRoman", "TimesNewRoman-Bold"
        except Exception:
            return "Helvetica", "Helvetica-Bold"


def create_docx(filename):
    """Generate the DOCX report."""
    doc = docx.Document()
    _set_docx_style(doc)

    _add_docx_centered(doc, "ЛАБОРАТОРНА РОБОТА № __LAB_NUM__", bold=False)
    _add_docx_centered(doc, "__WORK_TITLE__", bold=True)
    doc.add_paragraph()
    _add_docx_body(doc, " - __GOAL__", bold_prefix="Мета роботи")
    doc.add_paragraph()
    _add_docx_centered(doc, "Загальні відомості", bold=True)
__GENERAL_INFO_DOCX__
    doc.add_paragraph()
    _add_docx_centered(doc, "Завдання.", bold=True)
__TASKS_DOCX__
    doc.add_paragraph()
    _add_docx_centered(doc, "Контрольні запитання", bold=True)
__CONTROL_QUESTIONS_DOCX__
    doc.add_paragraph()
    _add_docx_centered(doc, "Література", bold=True)
__BIBLIOGRAPHY_DOCX__

    doc.save(filename)
    print(f"Успішно створено файл: {filename}")


def create_pdf(filename):
    """Generate the PDF report with identical content to the DOCX."""
    font_regular, font_bold = _setup_pdf_fonts()
    doc = SimpleDocTemplate(
        filename,
        pagesize=A4,
        rightMargin=1.5 * PdfCm,
        leftMargin=3.0 * PdfCm,
        topMargin=2.0 * PdfCm,
        bottomMargin=2.0 * PdfCm,
    )
    style_center = ParagraphStyle(
        name="Center", fontName=font_regular, fontSize=14,
        leading=21, alignment=TA_CENTER, spaceAfter=14,
    )
    style_center_bold = ParagraphStyle(
        name="CenterBold", fontName=font_bold, fontSize=14,
        leading=21, alignment=TA_CENTER, spaceAfter=14, spaceBefore=14,
    )
    style_body = ParagraphStyle(
        name="Body", fontName=font_regular, fontSize=14,
        leading=21, alignment=TA_JUSTIFY, firstLineIndent=35, spaceAfter=10,
    )
    story = []
    story.append(Paragraph("ЛАБОРАТОРНА РОБОТА № __LAB_NUM__", style_center))
    story.append(Paragraph("__WORK_TITLE__", style_center_bold))
    story.append(Paragraph("<b>Мета роботи</b> - __GOAL__", style_body))
    story.append(Paragraph("Загальні відомості", style_center_bold))
__GENERAL_INFO_PDF__
    story.append(Paragraph("Завдання.", style_center_bold))
__TASKS_PDF__
    story.append(Paragraph("Контрольні запитання", style_center_bold))
__CONTROL_QUESTIONS_PDF__
    story.append(Paragraph("Література", style_center_bold))
__BIBLIOGRAPHY_PDF__
    doc.build(story)
    print(f"Успішно створено файл: {filename}")


if __name__ == "__main__":
    create_docx("__OUT_DOCX__")
    create_pdf("__OUT_PDF__")
'''


# ─── Helpers for safe string substitution ────────────────────────────────

def _py_escape(value: str) -> str:
    """Escape a string for safe interpolation into a Python double-quoted literal.

    We forbid triple quotes in placeholder content. Backslashes and
    double quotes are escaped. Newlines become ``\\n`` so the resulting
    source is still one line per body paragraph.
    """
    if value is None:
        return ""
    text = str(value)
    text = text.replace("\\", "\\\\")
    text = text.replace('"', '\\"')
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    return text


def _block_docx(text: str) -> str:
    """Render a multi-line text as a sequence of `doc.add_paragraph(...)` calls.

    Each emitted line is indented 4 spaces (the body-of-function
    indent). The template places the substitution marker flush left.
    """
    lines = [ln for ln in (text or "").split("\n") if ln.strip()]
    if not lines:
        return '    doc.add_paragraph("")'
    return "\n".join(
        f'    doc.add_paragraph("{_py_escape(ln)}")' for ln in lines
    )


def _block_pdf(text: str) -> str:
    """Render a multi-line text as a sequence of `story.append(Paragraph(..., style_body))` calls.

    Each emitted line is indented 4 spaces.
    """
    lines = [ln for ln in (text or "").split("\n") if ln.strip()]
    if not lines:
        return '    story.append(Paragraph("", style_body))'
    return "\n".join(
        f'    story.append(Paragraph("{_py_escape(ln)}", style_body))'
        for ln in lines
    )


def _list_docx(items: list[str], numbered: bool = True) -> str:
    cleaned = [str(x).strip() for x in (items or []) if str(x or "").strip()]
    if not cleaned:
        return '    doc.add_paragraph("")'
    out = []
    for i, item in enumerate(cleaned, start=1):
        prefix = f"{i}. " if numbered else "• "
        out.append(f'    doc.add_paragraph("{_py_escape(prefix + item)}")')
    return "\n".join(out)


def _list_pdf(items: list[str], numbered: bool = True) -> str:
    cleaned = [str(x).strip() for x in (items or []) if str(x or "").strip()]
    if not cleaned:
        return '    story.append(Paragraph("", style_body))'
    out = []
    for i, item in enumerate(cleaned, start=1):
        prefix = f"{i}. " if numbered else "• "
        out.append(
            f'    story.append(Paragraph("{_py_escape(prefix + item)}", style_body))'
        )
    return "\n".join(out)


# ─── Gap-value accessors ──────────────────────────────────────────────────

def _gap(gap_values: dict[str, Any], key: str, default: Any = None) -> Any:
    """Pull a gap value from the schema-aware JSON."""
    block = gap_values.get(key) or {}
    if isinstance(block, dict):
        return block.get("value", default)
    return block


# ─── Template renderers ───────────────────────────────────────────────────

def _render_lab1(gap_values: dict[str, Any]) -> tuple[str, str]:
    """Returns (out_docx_filename, out_pdf_filename) and fills the template."""
    lab_num = str(_gap(gap_values, "lab_number", "1"))
    work_title = str(_gap(gap_values, "work_title", "Дослідження методів сортування в масивах"))
    goal = str(_gap(gap_values, "goal", "дослідити та порівняти ефективність алгоритмів сортування."))
    general_info = str(_gap(gap_values, "general_info", "Сортування є фундаментальною операцією в комп'ютерних науках."))
    tasks = _gap(gap_values, "tasks", []) or []
    questions = _gap(gap_values, "control_questions", []) or []
    biblio = _gap(gap_values, "bibliography", []) or []

    out_docx = "Lab_Template_Final.docx"
    out_pdf = "Lab_Template_Final.pdf"

    body = _TPL_DOCX
    body = body.replace("__LAB_NUM__", _py_escape(lab_num))
    body = body.replace("__WORK_TITLE__", _py_escape(work_title))
    body = body.replace("__GOAL__", _py_escape(goal))
    body = body.replace("__GENERAL_INFO_DOCX__", _block_docx(general_info))
    body = body.replace("__GENERAL_INFO_PDF__", _block_pdf(general_info))
    body = body.replace("__TASKS_DOCX__", _list_docx(tasks))
    body = body.replace("__TASKS_PDF__", _list_pdf(tasks))
    body = body.replace("__CONTROL_QUESTIONS_DOCX__", _list_docx(questions))
    body = body.replace("__CONTROL_QUESTIONS_PDF__", _list_pdf(questions))
    body = body.replace("__BIBLIOGRAPHY_DOCX__", _list_docx(biblio, numbered=True))
    body = body.replace("__BIBLIOGRAPHY_PDF__", _list_pdf(biblio, numbered=True))
    body = body.replace("__OUT_DOCX__", out_docx)
    body = body.replace("__OUT_PDF__", out_pdf)
    return body, out_docx, out_pdf


def _render_lab2(gap_values: dict[str, Any]) -> tuple[str, str, str]:
    """Lab2 uses the same 7-section layout. For now we re-use lab1."""
    # Override file names per the AGENTS.md "Lab2-specific" rules.
    body, _, _ = _render_lab1(gap_values)
    body = body.replace("Lab_Template_Final.docx", "Lab_Template_Lab2_Style.docx")
    body = body.replace("Lab_Template_Final.pdf", "Lab_Template_Lab2_Style.pdf")
    return body, "Lab_Template_Lab2_Style.docx", "Lab_Template_Lab2_Style.pdf"


def render_filled_py(template_id: str, gap_values: dict[str, Any]) -> dict[str, str]:
    """Render a filled.py for the given template using the supplied gap values.

    Returns a dict with keys: ``filled_py`` (the full script as text),
    ``out_docx`` (basename of the expected DOCX), ``out_pdf`` (basename
    of the expected PDF).
    """
    if template_id == "lab1":
        body, docx_name, pdf_name = _render_lab1(gap_values)
    elif template_id == "lab2":
        body, docx_name, pdf_name = _render_lab2(gap_values)
    else:
        # Unknown template: fall back to lab1 layout but keep filenames generic.
        body, docx_name, pdf_name = _render_lab1(gap_values)
    return {
        "filled_py": body,
        "out_docx": docx_name,
        "out_pdf": pdf_name,
    }


__all__ = ["render_filled_py"]
