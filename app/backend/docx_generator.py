"""
DSTU 3008:2015-compliant DOCX generator (python-docx).

Document structure follows the LMV_LabRob.pdf sample (newer style):
    Title page
    -> (optional) ВСТУП
    -> for each lab work:
        ЛАБОРАТОРНА РОБОТА №N       (centered, h1)
        Тема: <topic>               (centered, h2)
        Мета роботи - <objective>   (bold lead, body)
        Загальні відомості          (centered, h1)
        <theory paragraphs>
        Хід роботи                  (centered, h1)
        1. <step>                   (no «Крок N.» prefix duplication)
        Контрольні запитання        (centered, h1)
        1. <question>
        Завдання                    (centered, h1)
        1. <task>
        Варіанти                    (centered, h1)
        1. <variant>
        Зміст звіту                 (centered, h1)
        — <section>
        Література                  (centered, h1)
        1. <reference>
    -> (optional) СПИСОК РЕКОМЕНДОВАНИХ ДЖЕРЕЛ
"""

from __future__ import annotations

import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Inches, Pt

from .models import DocumentMetadata, LabGuidelinesContent


# ---------------------------------------------------------------------------
# Step-prefix normaliser (defensive — strips «Крок N.» that the AI sometimes
# prepends and the generator would duplicate as «1. Крок 1. …»).
# ---------------------------------------------------------------------------

_STEP_PREFIX_RE = re.compile(
    r"^\s*(?:крок|step|крок\s*№)\s*\d+\s*[\.\-:)]?\s*",
    flags=re.IGNORECASE,
)
_NUMERIC_PREFIX_RE = re.compile(r"^\s*\d+\s*[\.\-)]\s*")


def _strip_step_prefix(text: str) -> str:
    s = text.strip()
    changed = True
    while changed:
        changed = False
        if _STEP_PREFIX_RE.match(s):
            s = _STEP_PREFIX_RE.sub("", s, count=1).lstrip()
            changed = True
        elif _NUMERIC_PREFIX_RE.match(s):
            stripped = _NUMERIC_PREFIX_RE.sub("", s, count=1)
            if stripped != s and _STEP_PREFIX_RE.match(stripped):
                s = stripped.lstrip()
                changed = True
    return s


def add_page_number(run) -> None:
    """Inject Word fields for dynamic page numbering in headers."""
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


class DSTUDocxGenerator:
    def __init__(self) -> None:
        self.doc = Document()
        self._setup_styling()

    # ----- document-wide styling ------------------------------------------

    def _setup_styling(self) -> None:
        for section in self.doc.sections:
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(1.0)

            section.different_first_page_header_footer = True
            header = section.header
            p = header.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run()
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            add_page_number(run)

        normal = self.doc.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(14)
        normal.paragraph_format.line_spacing = 1.5
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.first_line_indent = Cm(1.25)

    # ----- building blocks ------------------------------------------------

    def _para(
        self,
        text: str = "",
        *,
        size: int = 14,
        bold: bool = False,
        italic: bool = False,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent: float | None = 1.25,
        space_after: float = 0,
        space_before: float = 0,
        line_spacing: float = 1.5,
    ):
        p = self.doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.first_line_indent = Cm(first_line_indent) if first_line_indent is not None else Cm(0)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.line_spacing = line_spacing
        if text:
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(size)
            run.bold = bold
            run.italic = italic
        return p

    def add_title_page(self, meta: DocumentMetadata) -> None:
        """Render the ДСТУ 3008:2015 title page.

        Layout: ministry / university / department at the top, the main
        title block just under, the authors on the right side. The city
        and year are pinned to the BOTTOM of the page (centred).
        python-docx has no absolute vertical positioning, so we push the
        city/year down with a small stack of empty spacer paragraphs.

        Long university / department / discipline / author strings are
        auto-shrunk (down to 8pt) so the title page never overflows.
        """
        def fit(text: str, *, size: int, bold: bool = False, italic: bool = False,
                align=WD_ALIGN_PARAGRAPH.CENTER, space_after: float = 0) -> int:
            """Write a paragraph, auto-shrinking `size` until the line
            fits inside the usable width. Returns the size actually used.
            """
            usable_cm = 21.0 - 3.0 - 1.0  # A4 - left - right (cm)
            # Rough char-width estimate: 0.5em per char in proportional fonts.
            # Use a slightly more generous estimate to be safe.
            approx_char_w = size * 0.42 / 28.35  # convert pt to cm, then × 0.42
            max_chars = usable_cm / approx_char_w if approx_char_w > 0 else len(text)
            current = size
            while current > 8 and len(text) * approx_char_w > usable_cm:
                current -= 1
                approx_char_w = current * 0.42 / 28.35
            self._para(
                text, size=current, bold=bold, italic=italic, align=align,
                first_line_indent=0, line_spacing=1.15, space_after=space_after,
            )
            return current

        # Top header block
        self._para("МІНІСТЕРСТВО ОСВІТИ І НАУКИ УКРАЇНИ", size=12, bold=True,
                   line_spacing=1.15, first_line_indent=0)
        fit(meta.university.upper(), size=12, bold=True)
        fit(f"Кафедра {meta.department}", size=12, space_after=24)

        # Main title block
        self._para("МЕТОДИЧНІ ВКАЗІВКИ", size=18, bold=True, space_before=12,
                   line_spacing=1.15, first_line_indent=0)
        self._para("до виконання лабораторних робіт з дисципліни", size=14,
                   line_spacing=1.15, first_line_indent=0)
        fit(f"«{meta.discipline}»", size=16, bold=True, space_after=60)

        # Authors (right-aligned, italic label + bold list of names)
        self._para(
            "Розробники:",
            size=12, italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT,
            space_before=12, line_spacing=1.15, first_line_indent=0,
        )
        fit(", ".join(meta.authors), size=12, bold=True,
            align=WD_ALIGN_PARAGRAPH.RIGHT)

        # Push city + year to the bottom of the page with empty paragraphs.
        # A4 has ~25.7cm of usable height (29.7 minus 2cm top + 2cm bottom
        # margins). The header + title + authors block above already uses
        # ~8-9cm. Each empty 12pt paragraph with 1.0 line spacing is
        # ~0.42cm tall, so 35-40 spacers fill the remaining space.
        spacers = 38
        for _ in range(spacers):
            self._para("", size=12, line_spacing=1.0, space_before=0, space_after=0,
                       first_line_indent=0)

        # City and year — pinned to the BOTTOM, centred.
        self._para(meta.city, size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
                   first_line_indent=0, line_spacing=1.15)
        self._para(str(meta.year), size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
                   first_line_indent=0, line_spacing=1.15, space_before=4)

        self.doc.add_page_break()

    def add_heading_h1(self, text: str) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(16)
        run.bold = True

    def add_heading_h1_centered(self, text: str) -> None:
        """Section heading used inside a lab work (centered, h1 sizing)."""
        self.add_heading_h1(text)

    def add_heading_h2(self, text: str) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.bold = True

    def add_heading_h2_centered(self, text: str) -> None:
        """Sub-section heading used inside a lab work (centered, h2 sizing)."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.bold = True

    def add_body_paragraph(self, text: str) -> None:
        self._para(text, size=14)

    def add_list(self, items, *, numbered: bool = True, bold_numbers: bool = True) -> None:
        for i, item in enumerate(items, 1):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.line_spacing = 1.5
            prefix = f"{i}. " if numbered else "— "
            run_p = p.add_run(prefix)
            run_p.font.name = "Times New Roman"
            run_p.font.size = Pt(14)
            if numbered and bold_numbers:
                run_p.bold = True
            run_t = p.add_run(item)
            run_t.font.name = "Times New Roman"
            run_t.font.size = Pt(14)

    def add_italic_tip(self, text: str) -> None:
        self._para(text, size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0)

    # ----- main entrypoint ------------------------------------------------

    def add_objective_lmv(self, text: str) -> None:
        """Bold 'Мета роботи' lead, dash separator (LMV style)."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.5
        run_b = p.add_run("Мета роботи ")
        run_b.font.name = "Times New Roman"
        run_b.font.size = Pt(14)
        run_b.bold = True
        run_d = p.add_run("— ")
        run_d.font.name = "Times New Roman"
        run_d.font.size = Pt(14)
        run_t = p.add_run(text)
        run_t.font.name = "Times New Roman"
        run_t.font.size = Pt(14)

    def add_bulleted_list(self, items: list[str]) -> None:
        """Bullet list with '—' prefix (LMV style for 'Зміст звіту')."""
        for item in items:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.line_spacing = 1.5
            run_p = p.add_run("— ")
            run_p.font.name = "Times New Roman"
            run_p.font.size = Pt(14)
            run_t = p.add_run(item)
            run_t.font.name = "Times New Roman"
            run_t.font.size = Pt(14)

    def add_numbered_steps(self, items: list[str]) -> None:
        """Numbered procedure list with «Крок N.» prefix auto-stripped."""
        for i, item in enumerate(items, 1):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.line_spacing = 1.5
            run_p = p.add_run(f"{i}. ")
            run_p.font.name = "Times New Roman"
            run_p.font.size = Pt(14)
            run_p.bold = True
            run_t = p.add_run(_strip_step_prefix(item))
            run_t.font.name = "Times New Roman"
            run_t.font.size = Pt(14)

    def generate(self, meta: DocumentMetadata, content: LabGuidelinesContent, filepath: str) -> None:
        # 1. Title page
        self.add_title_page(meta)

        # 2. Optional ВСТУП
        if content.introduction and content.introduction.strip():
            self.add_heading_h1("ВСТУП")
            for para in content.introduction.split("\n"):
                if para.strip():
                    self.add_body_paragraph(para.strip())
            self.doc.add_page_break()

        # 3. Lab works (LMV-style section order)
        for idx, lab in enumerate(content.lab_works, 1):
            # ЛАБОРАТОРНА РОБОТА №N (centered, h1)
            self.add_heading_h1_centered(f"ЛАБОРАТОРНА РОБОТА №{idx}")
            # Тема: <topic> (centered, h2)
            self.add_heading_h2_centered(f"Тема: {lab.topic}")
            # Мета роботи - <objective>
            self.add_objective_lmv(lab.objective)

            # Загальні відомості
            self.add_heading_h1_centered("Загальні відомості")
            for para in (lab.theory or "").split("\n"):
                if para.strip():
                    self.add_body_paragraph(para.strip())

            # Хід роботи
            if lab.procedure:
                self.add_heading_h1_centered("Хід роботи")
                self.add_numbered_steps(lab.procedure)

            # Контрольні запитання
            if lab.questions:
                self.add_heading_h1_centered("Контрольні запитання")
                self.add_list(lab.questions, numbered=True, bold_numbers=False)

            # Завдання
            if lab.tasks:
                self.add_heading_h1_centered("Завдання")
                self.add_list(lab.tasks, numbered=True, bold_numbers=False)

            # Варіанти
            if lab.variants:
                self.add_heading_h1_centered("Варіанти")
                self.add_list(lab.variants, numbered=True, bold_numbers=False)

            # Зміст звіту (bulleted)
            if lab.report_sections:
                self.add_heading_h1_centered("Зміст звіту")
                self.add_bulleted_list(lab.report_sections)

            # Література
            if lab.references:
                self.add_heading_h1_centered("Література")
                self.add_list(lab.references, numbered=True, bold_numbers=False)

            self.doc.add_page_break()

        # 4. Optional global references
        if content.references:
            self.add_heading_h1("СПИСОК РЕКОМЕНДОВАНИХ ДЖЕРЕЛ")
            self.add_list(content.references, numbered=True, bold_numbers=False)

        # 5. Save
        self.doc.save(filepath)
        print(f"Generated DOCX: {filepath}")
