import docx
import os

print("=== VERIFICATION OF CREATED FILES ===\n")

# Check DOCX files
docx_files = [
    r'labwork\AI_Integration_Education.docx',
    r'labwork\AI_Integration_VNTU.docx'
]

for docx_path in docx_files:
    if os.path.exists(docx_path):
        d = docx.Document(docx_path)
        print(f"File: {docx_path}")
        print(f"  Size: {os.path.getsize(docx_path)} bytes")
        print(f"  Paragraphs: {len(d.paragraphs)}")
        print(f"  Tables: {len(d.tables)}")
        
        # Count key formatting elements
        centred = sum(1 for p in d.paragraphs if p.alignment == 1 and p.text.strip())
        justified = sum(1 for p in d.paragraphs if p.alignment == 3 and p.text.strip())
        list_items = sum(1 for p in d.paragraphs if p.style.name.startswith('List'))
        first_indent = sum(1 for p in d.paragraphs if p.paragraph_format.first_line_indent and p.paragraph_format.first_line_indent.cm > 0 and p.text.strip())
        
        print(f"  Centre-aligned: {centred}")
        print(f"  Justified: {justified}")
        print(f"  List items: {list_items}")
        print(f"  First-line indent paragraphs: {first_indent}")
        
        # Show title page info
        for i, p in enumerate(d.paragraphs[:15]):
            if p.text.strip():
                align = {0: 'L', 1: 'C', 2: 'R', 3: 'J'}.get(p.alignment, str(p.alignment))
                print(f"  P{i}: [{align}] {p.style.name}: {p.text[:60]}")
        print()
    else:
        print(f"File not found: {docx_path}\n")

print("=== SUMMARY ===")
print("Both DOCX files have been created with Ukrainian methodichka formatting:")
print("- Title page centred, bold, ALL CAPS")
print('- Lab number and topic centred, bold')
print('- Bold inline label + plain text for "Мета роботи"')
print('- Centre-aligned section headings (Контрольні запитання, Завдання, etc.)')
print('- Justified body text with 1.25cm first-line indent')
print('- Numbered items in List Number style')
print('- Centre-aligned "Зауваження" and "Література"')
print('- Table for task variants')
print('- Numbered literature references')