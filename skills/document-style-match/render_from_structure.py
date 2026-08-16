"""Render a structure snapshot's blocks straight into a DOCX.

Usage: python render_from_structure.py <style_contract.json> <content_plan.json> <output.docx>

`content_plan.json` is the sample's `blocks` array from Step 1 (either
analyzer produces one) with each block's `"text"` replaced by the new
content — everything else (`align`/`bold`/`size_pt`/`indent_cm`) copied
through unchanged, because that is the sample's formatting, not its
content. This is what Step 3 in SKILL.md calls instead of hand-writing a
new python-docx script every run: two live sessions each wrote a fresh
~13KB generator from scratch (mismatched lambda signatures, wrong default
alignment) to do exactly the mechanical part this script now does once,
tested.

A title-size block (size_pt within 0.5pt of the contract's title_size_pt)
is centered regardless of its recorded align — the same convention
`analyze_pdf.py` applies when *measuring* it, kept consistent here so a
PDF-derived title block that had to be forced to "center" at analysis time
does not need its align field hand-corrected before rendering.

Not covered — write these by hand into the output DOCX with `python-docx`
directly (or extend `content_plan.json` and this script together, they are
meant to be edited as a pair): numbered lists that must restart at 1,
tables, images, footnotes. See `template_generate.py` for the tested
patterns (`new_list()`, `table()`) for those cases; this script only ever
emits plain paragraphs.
"""
import sys
import json

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH as A
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

_ALIGN = {"center": A.CENTER, "right": A.RIGHT, "justify": A.JUSTIFY, "left": A.LEFT}


def _opt_pt(value):
    return None if value is None else Pt(value)


def _apply_spacing(p, before, after) -> None:
    """Set paragraph spacing to the measured value, or pin it to zero.

    python-docx's default template is not neutral (Normal carries its own
    space_after), so leaving these unset would substitute Word's spacing
    for the sample's. Zero is the honest stand-in when the sample itself
    stores nothing: the gaps then come from spacer paragraphs, which is
    how the sample makes them.
    """
    p.paragraph_format.space_before = before if before is not None else Pt(0)
    p.paragraph_format.space_after = after if after is not None else Pt(0)


def _styled_paragraph(doc, style_name):
    """Add a paragraph in `style_name`, falling back to Normal if absent.

    A missing style must not abort a render — the plan may name a style
    the blank template does not carry.
    """
    if style_name and style_name != "Normal":
        try:
            return doc.add_paragraph(style=style_name)
        except KeyError:
            pass
    return doc.add_paragraph()


def _neutralise(run, font, size, bold, *, styled: bool) -> None:
    """Give the run the sample's font, and defeat the template's own look.

    Applying a real `Heading 1` is what gives the document an outline, but
    the blank template's version of it is 14pt blue Calibri Light, so a
    styled paragraph needs its colour forced back to black. Only styled
    ones: recolouring every run would overwrite deliberate colour a plain
    paragraph might carry.

    `bold=None` means the block did not state it and the style decides —
    writing `False` there cancelled the bold the sample's Heading styles
    supply, and every heading rendered unbolded.
    """
    run.font.name = font
    run.font.size = size
    if bold is not None:
        run.bold = bold
    if styled:
        run.font.color.rgb = RGBColor(0, 0, 0)


def _restart_numbering(doc):
    """A numbering instance that restarts at 1, or None if unavailable.

    Paragraphs sharing `style="List Number"` share the style's numbering,
    so a second list continues the first instead of restarting — measured
    live: a "Завдання" list rendered 4, 5, 6.
    """
    try:
        style_num_id = int(doc.styles["List Number"].element
                           .find(qn("w:pPr")).find(qn("w:numPr")).find(qn("w:numId"))
                           .get(qn("w:val")))
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        abstract_id = next(n.find(qn("w:abstractNumId")).get(qn("w:val"))
                           for n in numbering.findall(qn("w:num"))
                           if n.get(qn("w:numId")) == str(style_num_id))
        new_id = max(int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))) + 1
    except Exception:
        return None

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), abstract_id)
    num.append(ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return new_id


def _bind_numbering(p, num_id, level) -> None:
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(nid)
    pPr.append(numPr)


def _add_table(doc, block, font, body_size) -> None:
    rows = block.get("rows") or []
    if not rows:
        return
    cols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=cols)
    if block.get("style"):
        try:
            t.style = block["style"]
        except KeyError:
            pass
    widths = block.get("col_widths_cm") or []
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell = t.cell(ri, ci)
            cell.paragraphs[0].text = ""
            run = cell.paragraphs[0].add_run(row[ci] if ci < len(row) else "")
            run.font.name, run.font.size = font, body_size
            if ci < len(widths) and widths[ci]:
                cell.width = Cm(widths[ci])


def render(contract_path: str, content_plan_path: str, output_path: str) -> None:
    c = json.load(open(contract_path, encoding="utf-8"))
    blocks = json.load(open(content_plan_path, encoding="utf-8"))
    if isinstance(blocks, dict):
        blocks = blocks["blocks"]

    font = c.get("body_font") or "Times New Roman"
    body_size_pt = c.get("body_size_pt") or 14
    body_size = Pt(body_size_pt)
    title_size_pt = c.get("title_size_pt")
    indent = Cm(c.get("body_indent_cm", 1.25))
    line_spacing = c.get("line_spacing", 1.15)
    # `None` means the sample sets no spacing, and then neither may we.
    # Defaulting these to 0/4pt wrote 2pt and 4pt gaps into 1433
    # paragraphs of a sample that had `None` everywhere — spacing invented
    # by the renderer rather than measured. Word's own default is not
    # neutral either, so an unmeasured value is pinned to 0 and the
    # rhythm is left to the spacer paragraphs, exactly as in the sample.
    heading_sb = _opt_pt(c.get("heading_space_before_pt"))
    heading_sa = _opt_pt(c.get("heading_space_after_pt"))

    doc = docx.Document()
    doc.styles["Normal"].font.name = font
    doc.styles["Normal"].font.size = body_size
    m = c.get("margins_cm", {})
    for sec in doc.sections:
        sec.page_width, sec.page_height = Cm(c["page_width_cm"]), Cm(c["page_height_cm"])
        sec.top_margin, sec.bottom_margin = Cm(m.get("top_margin", 2.0)), Cm(m.get("bottom_margin", 2.0))
        sec.left_margin, sec.right_margin = Cm(m.get("left_margin", 2.5)), Cm(m.get("right_margin", 1.5))

    # One fresh numbering instance per run of consecutive list blocks that
    # share a `list_id`, so each list in the output starts again at 1.
    numbering_for: dict[str, int] = {}
    previous_list_id = None

    for b in blocks:
        if b.get("kind") == "table":
            _add_table(doc, b, font, body_size)
            previous_list_id = None
            continue

        # A spacer is an empty paragraph, and it is not decoration: in this
        # document convention it is the *only* thing producing vertical
        # rhythm (measured on a real coursework sample — 801 empty
        # paragraphs, and `space_before`/`space_after` None on every
        # paragraph and in every style). Skipping empty blocks, as this
        # loop used to, reproduced the text with none of the spacing.
        if b.get("kind") == "spacer" or not b.get("text", "").strip():
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = line_spacing
            _apply_spacing(p, _opt_pt(b.get("space_before_pt")) or heading_sb,
                           _opt_pt(b.get("space_after_pt")) or heading_sa)
            if b.get("page_break_before"):
                p.paragraph_format.page_break_before = True
            continue

        text = b["text"]
        size_pt = b.get("size_pt")
        # "Title-size text is centered" only means anything when the title
        # size is *distinct* from the body size. In a sample whose headings
        # are bold rather than larger, `title_size_pt == body_size_pt`, and
        # this override then matched every ordinary body block: measured on
        # a real coursework run, 26 blocks the plan marked `justify` were
        # all rendered centered, and the whole document came out centered.
        title_is_distinct = (title_size_pt is not None and body_size_pt is not None
                             and title_size_pt - body_size_pt > 0.5)
        is_title = (title_is_distinct and size_pt is not None
                    and abs(size_pt - title_size_pt) < 0.5)
        align = "center" if is_title else b.get("align")
        is_heading = align in ("center", "right")
        list_id = b.get("list_id")

        if list_id is not None:
            p = _styled_paragraph(doc, "List Number")
            if list_id != previous_list_id or list_id not in numbering_for:
                new_id = _restart_numbering(doc)
                if new_id is not None:
                    numbering_for[list_id] = new_id
            if list_id in numbering_for:
                _bind_numbering(p, numbering_for[list_id], b.get("list_level", 0))
        else:
            p = _styled_paragraph(doc, b.get("style_name"))
        previous_list_id = list_id

        styled = bool(b.get("style_name")) and b["style_name"] != "Normal"
        # Leaving alignment unset lets the style supply it; a plain
        # paragraph with nothing to inherit from still needs the
        # convention's justified body text.
        if align is not None:
            p.alignment = _ALIGN[align]
        elif not styled:
            p.alignment = A.JUSTIFY
        p.paragraph_format.line_spacing = line_spacing
        if b.get("page_break_before"):
            p.paragraph_format.page_break_before = True
        _apply_spacing(p, _opt_pt(b.get("space_before_pt")) or heading_sb,
                       _opt_pt(b.get("space_after_pt")) or heading_sa)
        if not is_heading and list_id is None and not styled:
            # `is not None`, not truthiness. A block whose indent was
            # *measured as zero* means "flush left", and reading 0.0 as
            # "unset" replaced it with the body default instead. Measured on
            # a live run: the model noticed the substitution and worked
            # around it by writing `indent_cm: 0.001` into every block of the
            # plan, which turned paragraph indentation off for the whole
            # document — and nothing downstream objected, because a
            # PDF-derived contract carries no `body_indent_cm` to check it
            # against. Only a genuinely absent key falls back now.
            indent_cm = b.get("indent_cm")
            p.paragraph_format.first_line_indent = (
                Cm(indent_cm) if indent_cm is not None else indent)

        _neutralise(p.add_run(text), font,
                    Pt(size_pt) if size_pt else body_size,
                    b.get("bold"), styled=styled)

    doc.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("Usage: python render_from_structure.py <style_contract.json> <content_plan.json> <output.docx>")
        raise SystemExit(2)
    render(sys.argv[1], sys.argv[2], sys.argv[3])
    print(f"Rendered: {sys.argv[3]}")
