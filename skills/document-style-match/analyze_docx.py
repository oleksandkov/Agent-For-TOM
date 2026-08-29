"""Measure a DOCX sample's style contract: page size, margins, fonts, spacing.

Usage: python analyze_docx.py <sample.docx> <contract.json>

Unlike a PDF sample, every number here is read directly from the section,
style and paragraph properties python-docx exposes — no bounding-box
guesswork needed. This matters for spacing specifically: measuring heading
space-before/after from a *PDF's* rendered line positions was tried and
does not generalize (Word merges adjacent paragraphs' before/after values
by taking the max, not the sum, and which one wins depends on what precedes
each heading — not recoverable from glyph positions alone, confirmed by
measuring three headings in the same real document and getting three
different apparent rules). A DOCX sample has no such ambiguity: these are
the paragraph's own stored properties.

`body_indent_cm` and `line_spacing` are the *mode* among justified
paragraphs (the common body-paragraph style), not the first one seen — a
single differently-formatted paragraph must not skew the contract.
`heading_space_before_pt`/`heading_space_after_pt` are the mode among
centered+bold paragraphs (headings), which can legitimately differ from
body spacing.

`blocks` is a structure snapshot: one entry per non-empty paragraph, in
document order, each carrying its own alignment/bold/size/indent — exact,
because these are the paragraph's own stored properties, not inferred from
glyph positions the way `analyze_pdf.py`'s `blocks` has to be. Capped at
`_BLOCK_LIMIT` paragraphs (comfortably more than one repeating unit in a
template that repeats the same section shape several times) so a long
document does not dump its entire body into the contract.
"""
import sys
import json
from collections import Counter

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH as ALIGN
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

EMU_PER_CM = 360000
MARGIN_KEYS = ("top_margin", "bottom_margin", "left_margin", "right_margin")
_ALIGN_NAME = {ALIGN.CENTER: "center", ALIGN.RIGHT: "right", ALIGN.JUSTIFY: "justify"}
_BLOCK_LIMIT = 400
# Sample text is carried only so a reader can tell blocks apart; the words
# themselves are replaced at generation time. Full text made a real
# coursework contract 79% sample prose and unreadable — measured.
_TEXT_PREVIEW_CHARS = 60
# A sample table can be hundreds of rows (a code listing laid out as one);
# a handful shows the shape, which is all a content plan needs.
_TABLE_PREVIEW_ROWS = 12


def _mode(counter: Counter, default=None):
    return counter.most_common(1)[0][0] if counter else default


def _page_break_before(p) -> bool:
    """True when Word starts this paragraph on a fresh page.

    Two mechanisms produce the same visual result and both must be caught:
    the paragraph's own `w:pageBreakBefore`, and a `<w:br w:type="page"/>`
    run closing the *previous* paragraph (the caller handles that half).
    Measured on a real coursework sample: 6 of the former, 19 of the
    latter — reading only one of the two finds under a quarter of them.
    """
    pPr = p._p.find(qn("w:pPr"))
    if pPr is None:
        return False
    el = pPr.find(qn("w:pageBreakBefore"))
    if el is None:
        return False
    return el.get(qn("w:val")) not in ("0", "false")


def _ends_with_page_break(p) -> bool:
    return bool(p._p.findall(".//" + qn("w:br") + '[@' + qn("w:type") + '="page"]'))


def _resolve_inherited(d):
    """Font name/size from `w:docDefaults`, which `styles["Normal"]` misses.

    A Word-authored document commonly leaves the Normal style's own font
    unset and inherits from `w:rPrDefault` instead — measured on a real
    coursework sample, where `styles["Normal"].font` was (None, None) while
    the document plainly rendered Times New Roman 14pt. Returning those
    Nones made the contract unusable and sent one session unzipping the
    .docx to read the XML by hand.
    """
    try:
        rpr = d.styles.element.find(qn("w:docDefaults")) \
                              .find(qn("w:rPrDefault")) \
                              .find(qn("w:rPr"))
    except AttributeError:
        return None, None
    if rpr is None:
        return None, None
    name = size = None
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is not None:
        name = fonts.get(qn("w:ascii")) or fonts.get(qn("w:hAnsi"))
    sz = rpr.find(qn("w:sz"))
    if sz is not None and sz.get(qn("w:val")):
        size = float(sz.get(qn("w:val"))) / 2  # half-points
    return name, size


def iter_body_items(document):
    """Paragraphs *and* tables, in the order they appear in the document.

    `document.paragraphs` silently skips tables, so a sample whose title
    page, table of contents and comparison table are tables (measured: 4
    of them in a real coursework) produced a block list with that content
    missing and the surrounding paragraphs in a misleading linear order.
    """
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _numpr_of(element):
    pPr = element.find(qn("w:pPr"))
    return pPr.find(qn("w:numPr")) if pPr is not None else None


def _list_info(p) -> tuple[str | None, int]:
    """(numbering id, level) when the paragraph is a real Word list item.

    Numbering can sit on the paragraph *or* on its style — a paragraph
    added as `style="List Number"` carries no `numPr` of its own, so
    reading only the paragraph found no lists at all. The style's own name
    is used as the identity in that case, which is enough for the renderer:
    it restarts numbering whenever a run of list blocks is interrupted.
    """
    numPr = _numpr_of(p._p)
    if numPr is None:
        style_numPr = _from_style_chain(p.style, lambda s: _numpr_of(s.element))
        if style_numPr is None:
            return None, 0
        return (p.style.name, 0)
    num_id = numPr.find(qn("w:numId"))
    ilvl = numPr.find(qn("w:ilvl"))
    return (num_id.get(qn("w:val")) if num_id is not None else None,
            int(ilvl.get(qn("w:val"))) if ilvl is not None else 0)


def _pt(value):
    return round(value.pt, 1) if value is not None else None


def _from_style_chain(style, get):
    """Walk `base_style` until one level answers, mimicking Word's lookup."""
    seen = 0
    while style is not None and seen < 10:
        value = get(style)
        if value is not None:
            return value
        style, seen = style.base_style, seen + 1
    return None


def _effective_align(p):
    """The alignment the sample actually renders with.

    Recording only the paragraph's own value loses the sample's intent:
    its Heading 1 paragraphs store `None` and take CENTER from a
    *customised* style. A blank template's Heading 1 has no alignment at
    all, so "inherit" would render them flush left. Resolving the value
    here and writing it explicitly reproduces the look on any template.
    """
    if p.alignment is not None:
        return p.alignment
    return _from_style_chain(p.style, lambda s: s.paragraph_format.alignment)


def _effective_bold(p, runs):
    if any(r.bold for r in runs):
        return True
    if runs and any(r.bold is False for r in runs):
        return False
    return bool(_from_style_chain(p.style, lambda s: s.font.bold))


def _table_block(t) -> dict:
    widths = []
    for col in t.columns:
        try:
            widths.append(round(col.width / EMU_PER_CM, 2) if col.width else None)
        except (AttributeError, TypeError, ValueError):
            widths.append(None)
    return {
        "kind": "table",
        "style": t.style.name if t.style is not None else None,
        "col_widths_cm": widths,
        "rows": [[c.text.strip() for c in row.cells] for row in t.rows],
    }


#: What a paragraph can carry that `edit_copy` copies but cannot edit.
_EMBEDDED_TAGS = {
    "math": "{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath",
    "images": "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline",
    "links": "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink",
}


def _embedded_counts(p) -> dict:
    """Formulas, images and hyperlinks this paragraph holds.

    Route A carries these for free and can edit none of them, and the block
    that holds one behaves differently from every other block: replacing its
    text leaves the equation in place, beside prose that no longer refers to
    it. That was invisible until after the copy. Measured: a session
    re-theming a numerical-integration paper to AI found 28 equations it
    could not remove through the edit file, so it built a doctored copy of
    the sample instead and fed that to `build` as the reference — the one
    thing rule 2 forbids, made unavoidable by the tool.
    """
    counts = {}
    for kind, tag in _EMBEDDED_TAGS.items():
        n = len(p._p.findall(f".//{tag}"))
        if n:
            counts[kind] = n
    return counts


def _blocks(document, *, limit: int = _BLOCK_LIMIT) -> list[dict]:
    """Every body item: paragraphs, tables, and the empty paragraphs too.

    Empty paragraphs are not noise in this document convention — they are
    the *only* thing producing vertical rhythm. Measured on a real
    coursework sample: 801 of 5497 paragraphs are empty, `space_before`
    and `space_after` are None on every paragraph *and* in every style
    definition, and headings sit between two of them. Dropping them (as
    this function used to) reproduced the text with none of the spacing.
    """
    blocks: list[dict] = []
    break_pending = False
    for item in iter_body_items(document):
        if isinstance(item, Table):
            blocks.append(_table_block(item))
            if len(blocks) >= limit:
                break
            continue

        p = item
        text = p.text.strip()
        pbb = _page_break_before(p) or break_pending
        break_pending = _ends_with_page_break(p)
        pf = p.paragraph_format
        if not text:
            block = {"kind": "spacer"}
        else:
            runs = [r for r in p.runs if r.text.strip()]
            indent = pf.first_line_indent
            # Alignment and bold are resolved through the style chain, not
            # read off the paragraph alone. The sample's Heading 1
            # paragraphs store alignment=None and bold=None and take CENTER
            # and bold from a customised style; recording the raw None made
            # the renderer emit them flush left and unbolded on a template
            # whose own Heading 1 supplies neither.
            align = _effective_align(p)
            block = {
                "text": text,
                "align": _ALIGN_NAME.get(align) if align is not None else None,
                "bold": _effective_bold(p, runs),
                "size_pt": round(runs[0].font.size.pt, 1) if runs and runs[0].font.size else None,
                "indent_cm": round(indent / EMU_PER_CM, 2) if indent and indent > 0 else 0.0,
            }
            # A heading is a *style*, not just centered bold text: carrying
            # it is what gives the output a navigable outline and makes a
            # table-of-contents field possible. Measured: a sample with 8
            # Heading 1 and 14 Heading 2 paragraphs produced an output with
            # none, every paragraph flattened to Normal.
            style_name = p.style.name if p.style is not None else None
            if style_name and style_name != "Normal":
                block["style_name"] = style_name
            list_id, list_level = _list_info(p)
            if list_id is not None:
                block["list_id"] = list_id
                block["list_level"] = list_level
        # Also on a spacer: a display equation is a paragraph with no text,
        # so the blocks that hold the sample's formulas are exactly the ones
        # recorded as spacers, and annotating only the prose would hide them.
        embedded = _embedded_counts(p)
        if embedded:
            block["embedded"] = embedded
        for key, value in (("space_before_pt", _pt(pf.space_before)),
                           ("space_after_pt", _pt(pf.space_after))):
            if value is not None:
                block[key] = value
        if pbb:
            block["page_break_before"] = True
        blocks.append(block)
        if len(blocks) >= limit:
            break
    return blocks


def style_signatures(blocks) -> list[dict]:
    """The distinct formatting combinations, with how often each occurs.

    120 blocks of a real coursework used 15 of these — the formatting
    information in a contract is small even when the document is not.
    """
    sig: Counter = Counter()
    for b in blocks:
        if b.get("kind") in ("spacer", "table"):
            continue
        sig[(b["align"], b["bold"], b["size_pt"], b["indent_cm"])] += 1
    return [{"align": a, "bold": bo, "size_pt": s, "indent_cm": i, "count": n}
            for (a, bo, s, i), n in sig.most_common()]


def split_contract(full: dict, total_blocks: int, *,
                   scan: dict | None = None) -> tuple[dict, dict]:
    """Separate the always-needed style numbers from the bulky block list.

    The style half is read every run and must stay small enough to read in
    one call; the structure half is read only when building a content
    plan. Keeping them in one file produced a 211 KB contract that one
    session spent 11 `read_file` calls failing to get through.

    Every block gets a `source_index`: its position in this list. That is
    what makes a content plan checkable — `check_plan` can only otherwise
    compare *counts*, and a plan written from scratch with a plausible
    number of blocks passes a count check. It also removes the arithmetic
    that cost two of three measured sessions their turn: the model was
    subtracting the JSON header's line count from `read_file`'s line numbers
    to guess the index, and got it wrong three times in a row.

    `scan` describes how much of the document was looked at, so `truncated`
    stops being the only signal. It read as "hit the block limit" on a
    47-page sample where the real reason was a 4-page scan window.
    """
    blocks = full.get("blocks", [])
    style = {k: v for k, v in full.items() if k != "blocks"}
    style["style_signatures"] = style_signatures(blocks)
    # `total_blocks` counts the document's own units — rendered rows for a
    # PDF, paragraphs for a DOCX. A PDF's blocks are now *paragraphs* built
    # from several rows each, so comparing the two directly reads as heavy
    # truncation when nothing was dropped: on the sample, 41 paragraphs
    # against 1467 rows. `rows_recorded` is what makes the two comparable.
    rows_recorded = sum(int(b.get("source_rows", 1)) for b in blocks)
    truncated = rows_recorded < total_blocks
    structure = {
        "total_blocks": total_blocks,
        "blocks_recorded": len(blocks),
        "rows_recorded": rows_recorded,
        "truncated": truncated,
        "blocks": [dict(_preview(b), source_index=i)
                   for i, b in enumerate(blocks)],
    }
    if truncated:
        structure["truncated_reason"] = _truncation_reason(len(blocks), scan)
    if scan:
        structure.update(scan)
    return style, structure


def _truncation_reason(recorded: int, scan: dict | None) -> str:
    """Why the structure is shorter than the document, in the reader's terms.

    Two different causes were reported identically. `block_limit` means the
    whole document was walked and the cap stopped it — the rest is more of
    the same. `page_window` means most of the document was never opened, and
    that is the one a caller has to act on: a sample whose later sections
    (Контрольні запитання, Завдання, Література) live on unscanned pages
    cannot be reproduced from this structure at all.
    """
    if recorded >= _BLOCK_LIMIT:
        return "block_limit"
    if scan and len(scan.get("pages_scanned") or []) < (scan.get("page_count") or 0):
        return "page_window"
    return "unknown"


def _clip(text: str) -> str:
    return text if len(text) <= _TEXT_PREVIEW_CHARS else text[:_TEXT_PREVIEW_CHARS] + "…"


def _preview(b: dict) -> dict:
    if b.get("kind") == "table":
        out = dict(b)
        rows = out.get("rows", [])
        out["rows"] = [[_clip(c) for c in row] for row in rows[:_TABLE_PREVIEW_ROWS]]
        if len(rows) > _TABLE_PREVIEW_ROWS:
            out["total_rows"] = len(rows)
        return out
    if "text" not in b:
        return b
    out = dict(b)
    t = out["text"]
    if len(t) > _TEXT_PREVIEW_CHARS:
        out["text"] = _clip(t)
        out["full_len"] = len(t)
    return out


def analyze(docx_path: str) -> dict:
    d = docx.Document(docx_path)
    section = d.sections[0]

    def emu_to_cm(emu: int) -> float:
        return round(emu / EMU_PER_CM, 2)

    normal_size = d.styles["Normal"].font.size
    inherited_font, inherited_size = _resolve_inherited(d)
    body_size = normal_size.pt if normal_size else inherited_size
    contract = {
        "page_width_cm": emu_to_cm(section.page_width),
        "page_height_cm": emu_to_cm(section.page_height),
        "margins_cm": {k: emu_to_cm(getattr(section, k)) for k in MARGIN_KEYS},
        "body_font": d.styles["Normal"].font.name or inherited_font,
        "body_size_pt": body_size,
    }

    # An *explicit* run size only wins if it is larger than the inherited
    # body size. Taking the largest explicit size unconditionally picked a
    # stray 10pt caption as the "title" of a document whose body is an
    # inherited 14pt — measured on a real coursework sample, and it then
    # made verify_docx.py fail a correct document.
    run_sizes = sorted(
        {r.font.size.pt for p in d.paragraphs for r in p.runs if r.font.size},
        reverse=True,
    )
    largest = run_sizes[0] if run_sizes else None
    if largest is not None and (body_size is None or largest > body_size):
        contract["title_size_pt"] = largest
    else:
        contract["title_size_pt"] = body_size or 12.0

    indents: Counter = Counter()
    line_spacings: Counter = Counter()
    heading_sb: Counter = Counter()
    heading_sa: Counter = Counter()

    for p in d.paragraphs:
        if not p.text.strip():
            continue
        is_heading = (p.alignment == ALIGN.CENTER
                      and any(r.bold for r in p.runs if r.text.strip()))
        pf = p.paragraph_format
        if is_heading:
            if pf.space_before is not None:
                heading_sb[round(pf.space_before.pt, 1)] += 1
            if pf.space_after is not None:
                heading_sa[round(pf.space_after.pt, 1)] += 1
        elif p.alignment == ALIGN.JUSTIFY:
            # Hanging-indent list items (dash/numbered) use a *negative*
            # first_line_indent and commonly outnumber plain body paragraphs
            # in a document full of questions and lists — including them
            # made the mode pick -0.63cm (a hanging indent) over 1.92cm (the
            # real body indent), confirmed against a real generated file
            # where they outnumbered true body paragraphs 25 to 7. A body
            # paragraph's first-line indent is always positive in this
            # convention, so excluding non-positive values removes them.
            if pf.first_line_indent is not None and pf.first_line_indent > 0:
                indents[round(pf.first_line_indent / EMU_PER_CM, 2)] += 1
            if pf.line_spacing is not None:
                line_spacings[round(pf.line_spacing, 2)] += 1

    contract["body_indent_cm"] = _mode(indents, 1.25)
    contract["line_spacing"] = _mode(line_spacings, 1.15)
    # `None` means "the sample sets nothing here", and the renderer must
    # then set nothing either. Defaulting these to 4.0 wrote 2pt/4pt gaps
    # into 1433 paragraphs of a sample that had `None` on every paragraph
    # and in every style definition — spacing invented by the tool, not
    # measured. The vertical rhythm comes from spacer paragraphs instead.
    contract["heading_space_before_pt"] = _mode(heading_sb, None)
    contract["heading_space_after_pt"] = _mode(heading_sa, None)
    contract["blocks"] = _blocks(d)
    contract["_total_blocks"] = sum(1 for _ in iter_body_items(d))
    return contract


def _structure_path(contract_path: str) -> str:
    if contract_path.endswith("_style_contract.json"):
        return contract_path[: -len("_style_contract.json")] + "_structure.json"
    if contract_path.endswith(".json"):
        return contract_path[: -len(".json")] + "_structure.json"
    return contract_path + "_structure.json"


# ensure_ascii=False everywhere: escaping Cyrillic to \uXXXX inflated a
# real contract from 50 KB to 211 KB — a 4.2x cost that exists only for
# non-English samples, which is every sample this skill is used on.
def _write(path: str, payload: dict) -> None:
    with open(path, "w", encoding="utf-8") as f:
        json.dump(payload, f, indent=2, ensure_ascii=False)


def _write_structure(path: str, structure: dict) -> None:
    """One block per line — pretty-printing each field cost 60% more.

    The blocks list is the whole point of this file and is read as a list,
    not field by field; `indent=2` spread each block over 8 lines for no
    reader benefit.
    """
    head = {k: v for k, v in structure.items() if k != "blocks"}
    with open(path, "w", encoding="utf-8") as f:
        f.write("{\n")
        for k, v in head.items():
            f.write(f"  {json.dumps(k)}: {json.dumps(v, ensure_ascii=False)},\n")
        f.write('  "blocks": [\n')
        lines = [f"    {json.dumps(b, ensure_ascii=False)}" for b in structure["blocks"]]
        f.write(",\n".join(lines))
        f.write("\n  ]\n}\n")


if __name__ == "__main__":
    if len(sys.argv) not in (3, 4):
        print("Usage: python analyze_docx.py <sample.docx> <contract.json> [structure.json]")
        raise SystemExit(2)
    full = analyze(sys.argv[1])
    total = full.pop("_total_blocks")
    style, structure = split_contract(full, total)
    contract_path = sys.argv[2]
    structure_path = sys.argv[3] if len(sys.argv) == 4 else _structure_path(contract_path)
    _write(contract_path, style)
    _write_structure(structure_path, structure)
    # Print the style half only. Printing the whole thing added a
    # 30,065-char tool result on top of the file that was just written.
    print(json.dumps(style, indent=2, ensure_ascii=False))
    print(f"\nstructure -> {structure_path} "
          f"({structure['blocks_recorded']} of {structure['total_blocks']} blocks"
          f"{', TRUNCATED' if structure['truncated'] else ''})")
