"""Reference implementation of every formatting pattern this skill needs.

Not meant to be run as-is — it has no real content. Read it for the parts
Step 3 in SKILL.md doesn't have room to inline: a second numbered list that
restarts at 1 instead of continuing, and a table. Adapt the parts you need
into your own generation script; keep the contract-driven values (FONT,
BODY, TITLE, INDENT, LS, HSB, HSA) exactly as shown — those are what make
the output match the measured sample instead of a guessed default.
"""
import json

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH as A
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

c = json.load(open(r"<name>_style_contract.json", encoding="utf-8"))
FONT = c.get("body_font") or "Times New Roman"
BODY, TITLE = Pt(c["body_size_pt"] or 14), Pt(c["title_size_pt"] or 20)
INDENT = Cm(c.get("body_indent_cm", 1.25))
LS = c.get("line_spacing", 1.15)
HSB = Pt(c.get("heading_space_before_pt", 0))
HSA = Pt(c.get("heading_space_after_pt", 4))

doc = docx.Document()
doc.styles["Normal"].font.name = FONT
doc.styles["Normal"].font.size = BODY
for sec in doc.sections:
    sec.page_width, sec.page_height = Cm(c["page_width_cm"]), Cm(c["page_height_cm"])
    m = c.get("margins_cm", {})
    sec.top_margin, sec.bottom_margin = Cm(m.get("top_margin", 2.0)), Cm(m.get("bottom_margin", 2.0))
    sec.left_margin, sec.right_margin = Cm(m.get("left_margin", 2.5)), Cm(m.get("right_margin", 1.5))


def para(text="", *, align=A.JUSTIFY, bold=False, indent=None, style=None, size=None, sb=Pt(0), sa=Pt(2)):
    p = doc.add_paragraph(style=style)
    p.alignment, p.paragraph_format.line_spacing = align, LS
    p.paragraph_format.space_before, p.paragraph_format.space_after = sb, sa
    if indent is not None:
        p.paragraph_format.first_line_indent = indent
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.font.name, r.font.size = FONT, (size or BODY)
    return p


body = lambda t: para(t, indent=INDENT)
centre = lambda t, size=None: para(t, align=A.CENTER, bold=True, size=size, sb=HSB, sa=HSA)


def lead(label, rest, indent=None):
    """'Мета роботи - текст' as ONE line, not a heading + new paragraph."""
    p = para(align=A.JUSTIFY, indent=indent or INDENT)
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.name, r1.font.size = FONT, BODY
    r2 = p.add_run(rest)
    r2.font.name, r2.font.size = FONT, BODY
    return p


# ── Numbered lists ──
# A plain `style="List Number"` paragraph inherits the *style's* numId, so a
# second numbered list continues the first instead of restarting at 1 —
# measured live: "Завдання" rendered 4, 5, 6 instead of 1, 2, 3. Call
# `new_list()` once per list; it returns an `item()` bound to a fresh
# numbering instance that starts at 1. Confirmed by rendering to PDF and
# reading the actual digits, not just inspecting the DOCX (python-docx
# cannot tell you what number Word will draw — only the render can).
def new_list():
    style_num_id = int(doc.styles["List Number"].element
                        .find(qn("w:pPr")).find(qn("w:numPr")).find(qn("w:numId"))
                        .get(qn("w:val")))
    numbering = doc.part.numbering_part.numbering_definitions._numbering
    abstract_id = next(n.find(qn("w:abstractNumId")).get(qn("w:val"))
                       for n in numbering.findall(qn("w:num"))
                       if n.get(qn("w:numId")) == str(style_num_id))
    new_id = max(int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))) + 1
    new_num = OxmlElement("w:num")
    new_num.set(qn("w:numId"), str(new_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), abstract_id)
    new_num.append(abs_ref)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    new_num.append(lvl_override)
    numbering.append(new_num)

    def item(text):
        p = para(text, style="List Number")
        pPr = p._p.get_or_add_pPr()
        numPr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(new_id))
        numPr.append(ilvl)
        numPr.append(num_id_el)
        pPr.append(numPr)
        return p

    return item


# ── Tables ──
def table(rows, *, header=True):
    """rows[0] is the header row if header=True. All cells same font/size."""
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = t.cell(ri, ci)
            cell.paragraphs[0].text = ""
            r = cell.paragraphs[0].add_run(str(cell_text))
            r.font.name, r.font.size = FONT, BODY
            r.bold = header and ri == 0
    return t


# ── Example usage ──
centre("<ACTUAL TITLE>", size=TITLE)
lead("<Мета роботи> - ", "<same line, not a new paragraph>")
body("<actual body text>")
centre("<Контрольні запитання>")
ask = new_list()
ask("<question 1>")
centre("<Завдання>")
task = new_list()  # a fresh new_list() call — reusing `ask` would continue its numbering
task("<task 1>")
table([["<Варіант>", "<Тема>"], ["1", "<...>"]])
doc.save(r"<output.docx>")
