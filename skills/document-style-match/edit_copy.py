"""Copy a DOCX sample and replace its words in place.

Usage:
  python edit_copy.py --list <sample.docx> [start] [count]
  python edit_copy.py <sample.docx> <edits.json> <output.docx>

For a DOCX sample this replaces the measure-and-rebuild route entirely.
Rebuilding reproduces only what the contract manages to describe, and
everything it misses is silently gone: measured on a real coursework run,
the rebuilt document had 0 of the sample's 4 tables, 0 of its 22
heading-styled paragraphs, 11 of its 116 spacer paragraphs and 57 of its
400 blocks — and every style check downstream still passed, because the
margins and fonts were right.

Copying cannot drift that way. Styles, section breaks, numbering, table
borders, column widths, headers and footers, and the vertical rhythm are
carried over because they are never re-derived — only the text changes.
It also costs far less: the edit file names the blocks that change
instead of restating every block in the document.

A PDF sample has no document to copy; that case still needs
`analyze_pdf.py` + `render_from_structure.py`.

## The edit file

Keys are block indices as printed by `--list` (and as positioned in
`<name>_structure.json`). Only listed blocks change; everything else is
kept exactly as the sample had it.

    {
      "3":  "New title for the work",
      "7":  ["Мета роботи - ", "нове формулювання мети"],
      "12": null,
      "31": {"rows": [["ВСТУП", "4"], ["1 РОЗДІЛ", "7"]]}
    }

  * string  — replace the paragraph's text, keeping its first run's
    formatting for the whole paragraph.
  * list    — one string per existing run, so a bold label followed by
    normal text stays a bold label followed by normal text.
  * null    — delete the paragraph or table.
  * object with "rows" — replace table cell text, keeping the table's
    own borders, widths and cell formatting.

Unlisted blocks that still hold sample text are counted and reported: on
this route the danger is not malformed output, it is the sample's own
sentences left behind in a document that is supposed to be new.
"""
import sys
import json
import shutil

import docx
from docx.table import Table

from analyze_docx import iter_body_items

#: Kept in step with `run.TODO_PREFIX`, and duplicated rather than imported:
#: `run.py` imports *this* module's siblings, so importing back would close a
#: cycle, and both scripts are run standalone as often as they are imported.
TODO_PREFIX = "<<TODO:"


def _clip(text: str, n: int = 70) -> str:
    text = " ".join(text.split())
    return text if len(text) <= n else text[:n] + "…"


def list_blocks(sample_path: str, start: int = 0, count: int = 60) -> None:
    d = docx.Document(sample_path)
    for i, item in enumerate(iter_body_items(d)):
        if i < start:
            continue
        if i >= start + count:
            print(f"… stopped at {start + count}; pass a start index to see more")
            break
        if isinstance(item, Table):
            print(f"{i:5}  [table {len(item.rows)}x{len(item.columns)}]  "
                  f"{_clip(' | '.join(c.text for c in item.rows[0].cells))}")
        elif item.text.strip():
            style = item.style.name if item.style is not None else ""
            tag = f"[{style}]" if style and style != "Normal" else ""
            print(f"{i:5}  {tag}{_clip(item.text)}")
        else:
            print(f"{i:5}  [spacer]")


def _coerce_runs(value, index) -> list[str]:
    """An edit value -> the texts to write into a paragraph's runs.

    This is the one line that produced a whole broken document. It used to be

        values = [value] if isinstance(value, str) else list(value)

    and `list()` on a **dict** yields its *keys*. Handed a block object like
    `{"text": "…", "align": "center", "bold": true, "size_pt": 16.0,
    "indent_cm": 0.0}` it wrote `text`, `align`, `bold`, `size_pt`,
    `indent_cm` into successive runs — so every paragraph of a real 52-block
    coursework rendered as `textalignboldsize_ptindent_cm`, and every spacer
    as `kind`. Nothing raised: `list()` succeeds on anything iterable, which
    is exactly why the wrong shape travelled all the way to a PDF.

    `list()` on an unknown type is not parsing. Anything that is not the
    documented shape is an error, said out loud, before a file is touched.
    """
    if isinstance(value, str):
        if value.lstrip().startswith(TODO_PREFIX):
            raise ValueError(
                f"block {index}: still holds the placeholder `{TODO_PREFIX}…` "
                f"that `run.py measure` wrote. Replace it with the text this "
                f"section should contain.")
        return [value]
    if isinstance(value, list):
        if all(isinstance(v, str) for v in value):
            return list(value)
        raise ValueError(
            f"block {index}: a list value must hold strings, one per run — "
            f"got {[type(v).__name__ for v in value][:4]}")
    if isinstance(value, dict):
        if "text" in value:
            raise ValueError(
                f"block {index}: the value is a content-plan block "
                f"({sorted(value)[:5]}…), but this route wants the new text "
                f"itself. Use the block's \"text\" as the value, or — if the "
                f"sample is a PDF — build a content plan and let run.py take "
                f"route B.")
        raise ValueError(
            f"block {index}: an object value is only meaningful for a table, "
            f"as {{\"rows\": [[…]]}} — got keys {sorted(value)[:5]}")
    raise ValueError(
        f"block {index}: expected a string, a list of strings, or null — "
        f"got {type(value).__name__}")


def validate_edits(edits: dict) -> None:
    """Check every value before anything is copied or written.

    Failing here costs nothing; failing after `shutil.copyfile` leaves a
    plausible-looking .docx on disk that a later step converts to PDF and
    reports as a deliverable. One measured run shipped exactly that.
    """
    problems = []
    for index, value in sorted(edits.items()):
        if value is None:
            continue
        if isinstance(value, dict) and "rows" in value:
            continue          # table; checked against the real item later
        try:
            _coerce_runs(value, index)
        except ValueError as e:
            problems.append(str(e))
    if problems:
        raise ValueError("the edit file does not match the documented shape:\n"
                         + "\n".join(f"  - {p}" for p in problems[:10])
                         + (f"\n  … and {len(problems) - 10} more"
                            if len(problems) > 10 else ""))


def _set_runs(p, value, index="?") -> None:
    values = _coerce_runs(value, index)
    runs = p.runs
    if not runs:
        for v in values:
            p.add_run(v)
        return
    for i, run in enumerate(runs):
        run.text = values[i] if i < len(values) else ""
    if len(values) > len(runs):
        last = runs[-1]
        for v in values[len(runs):]:
            new = p.add_run(v)
            new.bold, new.italic, new.underline = last.bold, last.italic, last.underline
            new.font.name, new.font.size = last.font.name, last.font.size


def _set_table(table, rows) -> None:
    for ri, row in enumerate(rows):
        if ri >= len(table.rows):
            break
        for ci, text in enumerate(row):
            if ci >= len(table.columns):
                break
            cell = table.cell(ri, ci)
            para = cell.paragraphs[0]
            if para.runs:
                _set_runs(para, text, f"table r{ri}c{ci}")
            else:
                para.add_run(text)


def apply_edits(sample_path: str, edits_path: str, output_path: str) -> dict:
    edits = json.load(open(edits_path, encoding="utf-8"))
    if isinstance(edits, dict) and "replacements" in edits:
        edits = edits["replacements"]
    if not isinstance(edits, dict):
        raise ValueError(
            f"{edits_path} holds a {type(edits).__name__}, but this route "
            f"wants an object keyed by block index: "
            f'{{"0": "new text", "3": null}}. A list of blocks is a *content '
            f"plan* — that is route B, for a PDF sample.")
    edits = {int(k): v for k, v in edits.items()}
    validate_edits(edits)          # before anything is written. See above.

    shutil.copyfile(sample_path, output_path)
    d = docx.Document(output_path)

    replaced = deleted = 0
    to_delete = []
    items = list(iter_body_items(d))
    for i, item in enumerate(items):
        if i not in edits:
            continue
        value = edits[i]
        if value is None:
            to_delete.append(item)
            deleted += 1
        elif isinstance(item, Table):
            if isinstance(value, dict) and "rows" in value:
                _set_table(item, value["rows"])
                replaced += 1
        else:
            _set_runs(item, value, i)
            replaced += 1

    for item in to_delete:
        element = item._tbl if isinstance(item, Table) else item._p
        element.getparent().remove(element)

    d.save(output_path)

    # What still carries the sample's words? Untouched blocks are the point
    # of this route — the title page's university name should survive — but
    # untouched *body prose* is the sample's content shipped as if it were
    # new, which is the one failure copying makes easy.
    sample = docx.Document(sample_path)
    untouched = 0
    for i, item in enumerate(iter_body_items(sample)):
        if i in edits or isinstance(item, Table):
            continue
        if len(item.text.strip()) > 120:
            untouched += 1
    return {"replaced": replaced, "deleted": deleted,
            "long_paragraphs_left_as_sample": untouched,
            "lost": _lost_embedded(sample, docx.Document(output_path))}


#: XML tags for the things this route carries but cannot edit, and their
#: names in a report.
_EMBEDDED = {
    "formulas": "{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath",
    "images": "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline",
    "hyperlinks": "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink",
}


def _count_embedded(document) -> dict:
    body = document.element.body
    counts = {name: len(body.findall(f".//{tag}"))
              for name, tag in _EMBEDDED.items()}
    counts["tables"] = len(document.tables)
    return counts


def _lost_embedded(sample, output) -> dict:
    """What the sample had and the output does not.

    Copying carries formulas, images, hyperlinks and tables for free — that
    is the whole argument for this route. What it cannot do is *edit* them,
    and a model that needs an equation changed has one obvious way out.
    Measured on a live run against a paper whose substance is its equations:
    `oMath elements removed: 33`, the output shipped with **0 of the
    sample's 28 formulas**, and every check still passed, because a formula
    is not a style field.

    Deleting them can be right — the user may have asked for prose. It must
    just never be silent.
    """
    before, after = _count_embedded(sample), _count_embedded(output)
    return {k: (before[k], after[k]) for k in before if after[k] < before[k]}


if __name__ == "__main__":
    if len(sys.argv) >= 3 and sys.argv[1] == "--list":
        start = int(sys.argv[3]) if len(sys.argv) > 3 else 0
        count = int(sys.argv[4]) if len(sys.argv) > 4 else 60
        list_blocks(sys.argv[2], start, count)
        raise SystemExit(0)
    if len(sys.argv) != 4:
        print(__doc__.split("## The edit file")[0].strip())
        raise SystemExit(2)

    stats = apply_edits(sys.argv[1], sys.argv[2], sys.argv[3])
    print(f"Wrote {sys.argv[3]}: {stats['replaced']} blocks replaced, "
          f"{stats['deleted']} deleted.")

    failed = False
    lost = stats["lost"]
    if lost:
        print("LOST FROM THE SAMPLE:")
        for name, (before, after) in sorted(lost.items()):
            print(f" - {name}: {before} in the sample, {after} in the output")
        print("These are carried by copying and cannot be edited by this "
              "script. If removing them was deliberate, say so in your final "
              "answer; otherwise restore them before shipping.")
        failed = True

    left = stats["long_paragraphs_left_as_sample"]
    if left:
        print(f"WARNING: {left} long paragraphs still hold the sample's own text. "
              f"Replace or delete them — run --list to find them.")
        failed = True
    if failed:
        raise SystemExit(1)
