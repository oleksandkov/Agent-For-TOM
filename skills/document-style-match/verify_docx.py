"""Diff a generated DOCX against its style contract. Not a self-check —
"some paragraphs are centered" would still pass a page-size-wrong document.

Usage: python verify_docx.py <contract.json> <output.docx>

Exits 0 and prints "Matches contract." when everything is within tolerance;
exits 1 and prints CONTRACT MISMATCH with the offending fields otherwise.
A mismatch means the generation script did not actually read the contract
for that field — fix that script and regenerate, do not report success.

Re-measures the output with the *same* `analyze()` used on the sample
(imported, not reimplemented) and diffs every field it returns — including
spacing/indent, not just page size and body font size — so a script that
read the contract for margins but hardcoded `line_spacing=1.15` regardless
of what the sample actually used gets caught here too.
"""
import os
import sys
import json

from analyze_docx import analyze, style_signatures

CM_TOLERANCE = 0.05
PT_TOLERANCE = 0.5

#: A formatting combination this share of the document is structural — it is
#: how a section heading, a body paragraph or a numbered item is set, and the
#: output is expected to use it too. Below this, one block happened to be
#: formatted that way and reproducing it is not required.
PROMINENT_SHARE = 0.05

#: How far a structural signature's share may fall in the output before it
#: counts as lost. Generous, because the output is a different document: the
#: sample's five centered headings may legitimately become four.
UNDERUSE_RATIO = 0.4

#: …and it has to be at least this many paragraphs, whatever the share says.
#: A share alone makes every single paragraph "structural" in a short
#: document — in an eleven-block one, a lone 9pt caption is 9% and got
#: reported as an invented formatting. Two is the smallest number that can be
#: a pattern rather than an instance.
PROMINENT_MIN_BLOCKS = 2


def _flatten(contract: dict) -> dict:
    flat = dict(contract)
    flat.update(flat.pop("margins_cm", {}))
    return flat


def verify_rhythm(structure_path: str, docx_path: str) -> list[str]:
    """Check the vertical rhythm the style fields cannot express.

    A document can match every number in the contract and still render as
    one unbroken wall of text: measured on two real outputs that both
    reported "Matches contract." while carrying zero spacer paragraphs and
    zero page breaks against a sample with 801 and 25. Spacing lives in
    the block list, so it needs the structure file, not the style half.
    """
    import docx
    from docx.oxml.ns import qn

    want = json.load(open(structure_path, encoding="utf-8"))
    blocks = want.get("blocks", [])
    want_spacers = sum(1 for b in blocks if b.get("kind") == "spacer")
    want_breaks = sum(1 for b in blocks if b.get("page_break_before"))

    d = docx.Document(docx_path)
    got_spacers = sum(1 for p in d.paragraphs if not p.text.strip())
    got_breaks = 0
    for p in d.paragraphs:
        pPr = p._p.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:pageBreakBefore")) is not None:
            got_breaks += 1
        got_breaks += len(p._p.findall(".//" + qn("w:br") + '[@' + qn("w:type") + '="page"]'))

    problems = []
    # Proportional, not exact: the generated document legitimately has a
    # different number of paragraphs than the sample. Losing spacing
    # altogether is the failure worth catching, not being a few off.
    if want_spacers and got_spacers < max(1, want_spacers * 0.3):
        problems.append(f"spacer paragraphs: got {got_spacers}, sample has {want_spacers} "
                        f"— vertical gaps between sections are missing")
    if want_breaks and got_breaks == 0:
        problems.append(f"page breaks: got 0, sample has {want_breaks} "
                        f"— sections that should start on a new page do not")
    return problems


def _triple(sig: dict) -> tuple:
    """A signature reduced to what must survive into the output.

    `align: None` is python-docx for "inherited", which renders left — and a
    heading that lost its centering arrives here as None, not as "left", so
    normalising is what makes the two comparable at all. Indent is bucketed to
    a millimetre because the renderer writes the contract's own number back
    and an exact float compare would fail on rounding, not on layout.
    """
    align = sig.get("align") or "left"
    indent = sig.get("indent_cm")
    return (str(align), bool(sig.get("bold")), sig.get("size_pt"),
            round(indent, 1) if isinstance(indent, (int, float)) else None)


def _describe(triple: tuple) -> str:
    align, bold, size, indent = triple
    return (f"{align}{'+bold' if bold else ''} at {size}pt"
            f"{f', indent {indent}cm' if indent else ''}")


def _profile(signatures: list[dict]) -> dict[tuple, tuple[float, int]]:
    """Each formatting combination's (share, count) over the real paragraphs.

    Shares carry the comparison: the output is a different document about a
    different topic and is entitled to a different length, and a count
    comparison called a correct 63-paragraph rebuild of a 41-paragraph sample
    "too many headings". Counts are kept alongside only to floor it — see
    `PROMINENT_MIN_BLOCKS`.
    """
    total = sum(s.get("count", 0) for s in signatures)
    if not total:
        return {}
    merged: dict[tuple, int] = {}
    for s in signatures:
        merged[_triple(s)] = merged.get(_triple(s), 0) + s.get("count", 0)
    return {t: (n / total, n) for t, n in merged.items()}


def verify_signatures(contract: dict, blocks: list) -> list[str]:
    """Did the output keep the sample's *formatting vocabulary*?

    Every other check here compares a number that is the same everywhere in
    the document — page size, body font, line spacing. None of them can see
    the thing a reader notices first, which is a heading that stopped being
    centered, and one measured run proves the gap is not theoretical: a
    rebuild set `Загальні відомості` flush left, dropped bold from the topic
    line and lost the 20pt title style altogether, and reported success
    because every scalar field it *did* set was right.

    Measured on that pair, sample against output:

        center+bold 20.0pt   7.3%  ->   absent   (title style gone)
        center+bold 14.0pt  12.2%  ->     1.6%   (headings no longer centered)
        justify 14.0pt/1.73  19.5% ->   absent   (numbered items lost indent)
        left+bold 14.0pt     absent ->    6.3%   (what they became instead)

    against a correct rebuild of the same sample, whose eight signatures all
    landed within about one percentage point. So the check is symmetric: a
    structural signature of the sample must survive into the output, and the
    output may not invent a prominent one the sample never uses. Anything
    rarer than `PROMINENT_SHARE` is left alone — that is one block formatted
    oddly, not a pattern.
    """
    want = _profile(contract.get("style_signatures") or [])
    got = _profile(style_signatures(blocks))
    if not want or not got:
        return []

    problems = []
    for triple, (share, count) in sorted(want.items(), key=lambda kv: -kv[1][0]):
        if share < PROMINENT_SHARE or count < PROMINENT_MIN_BLOCKS:
            continue
        mine = got.get(triple, (0.0, 0))[0]
        if mine >= share * UNDERUSE_RATIO:
            continue
        # Name what the output used instead. "center+bold 14pt is missing"
        # sends you looking for a lost paragraph; "the output sets those
        # left+bold" names the line to change. Candidates that kept the weight
        # rank first — a heading that lost its centering is still bold, and
        # pointing at the body paragraphs instead answered the wrong question.
        instead = sorted(
            (t for t in got if t[2] == triple[2] and t != triple),
            key=lambda t: (t[1] != triple[1], -got[t][0]))
        alternative = (f" — the output sets {_describe(instead[0])} instead"
                       if instead else "")
        problems.append(
            f"formatting lost: the sample sets {_describe(triple)} for "
            f"{share:.0%} of its paragraphs, the output for {mine:.0%}"
            f"{alternative}")

    for triple, (share, count) in sorted(got.items(), key=lambda kv: -kv[1][0]):
        if share < PROMINENT_SHARE or count < PROMINENT_MIN_BLOCKS:
            continue
        if triple in want:
            continue
        problems.append(
            f"formatting invented: {count} of the output's paragraphs "
            f"({share:.0%}) are {_describe(triple)}, which the sample never "
            f"uses")
    return problems


#: Measured off the output but not comparable to anything — reporting them as
#: unchecked would be noise, not information.
_NOT_A_STYLE_FIELD = frozenset({"blocks", "style_signatures"})


def verify(contract_path: str, docx_path: str) -> tuple[list[str], list[str]]:
    """Returns (problems, unchecked).

    `unchecked` is the second half of the answer, and it used to be silent.
    A field the contract does not carry is skipped here — correctly, there is
    nothing to compare against — but the caller then printed an unqualified
    "Matches contract." Measured live: a PDF-derived contract carried neither
    `line_spacing` nor `body_indent_cm`, the generated document had single
    spacing and no paragraph indent against a sample set to 1.5 and 1.25cm,
    and this said it matched. Both analyzers now measure both fields, so the
    gap should not recur — but a pass that names what it could not check is
    the thing that would have caught it, and it stays either way.
    """
    contract = _flatten(json.load(open(contract_path, encoding="utf-8")))
    measured = _flatten(analyze(docx_path))

    problems, unchecked = [], []
    for key, actual in measured.items():
        expected = contract.get(key)
        if key in _NOT_A_STYLE_FIELD or key.startswith("_"):
            # blocks carries the sample's own text for reference — the
            # generated document is expected to have different content, so
            # diffing it here would always "fail" on the one field that is
            # supposed to change.
            continue
        if expected is None:
            # Present-with-null and absent are different answers. A DOCX
            # sample that sets no heading spacing anywhere is *measured* as
            # having none, and the contract records that as null — there is
            # nothing to compare and nothing missing either. A field the
            # analyzer never looked at is simply not there, and that is the
            # one worth naming.
            if key not in contract and actual is not None:
                unchecked.append(key)
            continue
        if isinstance(expected, str):
            if str(actual) != expected:
                problems.append(f"{key}: got {actual!r}, contract {expected!r}")
            continue
        if actual is None:
            problems.append(f"{key}: not set in the output, contract {expected}")
            continue
        tolerance = PT_TOLERANCE if key.endswith("_pt") else CM_TOLERANCE
        if abs(actual - expected) > tolerance:
            problems.append(f"{key}: got {actual}, contract {expected}")

    # Runs last so the scalar mismatches — page size, fonts, spacing — read
    # first. They are usually the cause, and a document built on Letter paper
    # has bigger problems than its heading alignment.
    problems += verify_signatures(contract, measured.get("blocks", []))
    if not contract.get("style_signatures"):
        unchecked.append("style_signatures")
    return problems, unchecked


if __name__ == "__main__":
    if len(sys.argv) not in (3, 4):
        print("Usage: python verify_docx.py <contract.json> <output.docx> [structure.json]")
        raise SystemExit(2)
    problems, unchecked = verify(sys.argv[1], sys.argv[2])
    # The structure file is optional so an existing 2-argument call still
    # works, but without it the rhythm check cannot run — say so rather
    # than printing an unqualified pass.
    rhythm_checked = False
    structure_path = sys.argv[3] if len(sys.argv) == 4 else None
    if structure_path is None:
        guess = sys.argv[1]
        for suffix in ("_style_contract.json", ".json"):
            if guess.endswith(suffix):
                candidate = guess[: -len(suffix)] + "_structure.json"
                if os.path.exists(candidate):
                    structure_path = candidate
                break
    if structure_path and os.path.exists(structure_path):
        problems += verify_rhythm(structure_path, sys.argv[2])
        rhythm_checked = True

    if unchecked:
        print("NOT checked (the contract carries no value for these): "
              + ", ".join(sorted(unchecked)))
        print("  Re-measure the sample with the current analyzer, or compare "
              "these by eye on the rendered page — a pass below says nothing "
              "about them.\n")

    if problems:
        print("CONTRACT MISMATCH:")
        for p in problems:
            print(" -", p)
        raise SystemExit(1)
    verdict = ("Matches contract." if rhythm_checked
               else "Matches contract. (no structure file — spacing/page breaks NOT checked)")
    if unchecked:
        verdict += f" ({len(unchecked)} field(s) NOT checked — see above)"
    print(verdict)
